#!/usr/bin/env python3
"""
Comprehensive Export System for Academic Agent

This module provides a unified export system that supports multiple formats (PDF, HTML, DOCX)
with consolidated image handling, local reference resolution, template-based styling,
batch processing, and quality validation.

Key Features:
- Multiple export formats (PDF, HTML, DOCX)
- Consolidated image handling with optimized sizing
- Local reference resolution and linking
- Template-based export with customizable styling
- Batch export capabilities
- Integration with study notes generator
- Export quality validation and optimization
"""

import os
import json
import re
import time
import asyncio
import subprocess
import tempfile
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple, Union
from dataclasses import dataclass, asdict
from datetime import datetime
from abc import ABC, abstractmethod
import hashlib
import base64
from urllib.parse import quote
import zipfile

from dotenv import load_dotenv
from smolagents import Tool
import litellm

# Load environment variables
load_dotenv()

@dataclass
class ExportConfig:
    """Configuration for export operations"""
    output_format: str  # 'pdf', 'html', 'docx', 'all'
    template_name: str = 'academic'
    image_sizing: str = 'medium'  # 'small', 'medium', 'large', 'original'
    include_diagrams: bool = True
    diagram_format: str = 'png'  # 'png', 'svg', 'text'
    quality_level: str = 'high'  # 'low', 'medium', 'high'
    consolidate_images: bool = True
    embed_images: bool = True
    resolve_references: bool = True
    optimize_for_print: bool = True
    custom_css: Optional[str] = None
    metadata: Dict[str, Any] = None

@dataclass
class ExportResult:
    """Result of an export operation"""
    success: bool
    output_files: List[str]
    format_type: str
    file_size_mb: float
    processing_time: float
    validation_score: float
    errors: List[str]
    warnings: List[str]
    metadata: Dict[str, Any]

@dataclass
class ImageConsolidationResult:
    """Result of image consolidation"""
    consolidated_images: Dict[str, str]  # original_path -> consolidated_path
    total_size_mb: float
    optimization_ratio: float
    errors: List[str]

class ImageConsolidator:
    """Handles image consolidation and optimization"""
    
    def __init__(self, base_dir: Path):
        self.base_dir = base_dir
        self.consolidated_dir = base_dir / "export" / "images"
        self.consolidated_dir.mkdir(parents=True, exist_ok=True)
        
    def consolidate_images(self, content_paths: List[Path], 
                          sizing: str = 'medium') -> ImageConsolidationResult:
        """Consolidate and optimize images from multiple content sources"""
        
        consolidated_images = {}
        total_original_size = 0
        total_consolidated_size = 0
        errors = []
        
        # Size mappings
        size_configs = {
            'small': {'max_width': 400, 'quality': 70},
            'medium': {'max_width': 600, 'quality': 80},
            'large': {'max_width': 800, 'quality': 85},
            'original': {'max_width': None, 'quality': 90}
        }
        
        size_config = size_configs.get(sizing, size_configs['medium'])
        
        # Find all images in content
        image_paths = self._find_all_images(content_paths)
        
        for img_path in image_paths:
            try:
                # Calculate hash for unique naming
                with open(img_path, 'rb') as f:
                    content_hash = hashlib.md5(f.read()).hexdigest()[:8]
                
                # Generate consolidated filename
                consolidated_name = f"{content_hash}_{img_path.stem}{img_path.suffix}"
                consolidated_path = self.consolidated_dir / consolidated_name
                
                # Track original size
                total_original_size += img_path.stat().st_size
                
                # Optimize and copy image
                if self._optimize_image(img_path, consolidated_path, size_config):
                    consolidated_images[str(img_path)] = str(consolidated_path)
                    total_consolidated_size += consolidated_path.stat().st_size
                else:
                    # Fallback: copy original
                    shutil.copy2(img_path, consolidated_path)
                    consolidated_images[str(img_path)] = str(consolidated_path)
                    total_consolidated_size += consolidated_path.stat().st_size
                    
            except Exception as e:
                errors.append(f"Failed to consolidate {img_path}: {str(e)}")
        
        optimization_ratio = (
            (total_original_size - total_consolidated_size) / total_original_size
            if total_original_size > 0 else 0
        )
        
        return ImageConsolidationResult(
            consolidated_images=consolidated_images,
            total_size_mb=total_consolidated_size / (1024 * 1024),
            optimization_ratio=optimization_ratio,
            errors=errors
        )
    
    def _find_all_images(self, content_paths: List[Path]) -> List[Path]:
        """Find all image files referenced in content"""
        image_paths = set()
        image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.svg', '.bmp', '.webp'}
        
        for content_path in content_paths:
            if content_path.is_file():
                # Check content for image references
                try:
                    with open(content_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        
                    # Find markdown image references
                    img_refs = re.findall(r'!\[[^\]]*\]\(([^)]+)\)', content)
                    
                    for ref in img_refs:
                        if not ref.startswith(('http://', 'https://', 'data:')):
                            # Resolve relative path
                            img_path = content_path.parent / ref
                            if img_path.exists() and img_path.suffix.lower() in image_extensions:
                                image_paths.add(img_path.resolve())
                                
                except Exception:
                    pass
                    
            elif content_path.is_dir():
                # Scan directory for images
                for img_path in content_path.rglob('*'):
                    if img_path.suffix.lower() in image_extensions:
                        image_paths.add(img_path)
        
        return list(image_paths)
    
    def _optimize_image(self, source_path: Path, target_path: Path, 
                       config: Dict[str, Any]) -> bool:
        """Optimize an image file"""
        try:
            # Try using Pillow for optimization
            from PIL import Image
            
            with Image.open(source_path) as img:
                # Convert to RGB if necessary
                if img.mode in ('RGBA', 'LA', 'P'):
                    # Create white background
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                    img = background
                
                # Resize if needed
                if config['max_width'] and img.width > config['max_width']:
                    ratio = config['max_width'] / img.width
                    new_height = int(img.height * ratio)
                    img = img.resize((config['max_width'], new_height), Image.Resampling.LANCZOS)
                
                # Save with optimization
                save_kwargs = {
                    'quality': config['quality'],
                    'optimize': True
                }
                
                if target_path.suffix.lower() == '.png':
                    img.save(target_path, 'PNG', optimize=True)
                else:
                    img.save(target_path, 'JPEG', **save_kwargs)
                
            return True
            
        except ImportError:
            # Fallback: just copy if Pillow not available
            return False
        except Exception:
            return False

class ReferenceResolver:
    """Resolves local references and creates proper linking"""
    
    def __init__(self, base_dir: Path):
        self.base_dir = base_dir
        self.reference_map = {}
        
    def resolve_references(self, content: str, source_path: Path,
                          consolidated_images: Dict[str, str]) -> str:
        """Resolve all references in content"""
        
        # Resolve image references
        content = self._resolve_image_references(content, source_path, consolidated_images)
        
        # Resolve internal document references
        content = self._resolve_internal_references(content, source_path)
        
        # Resolve external file references
        content = self._resolve_file_references(content, source_path)
        
        return content
    
    def _resolve_image_references(self, content: str, source_path: Path,
                                 consolidated_images: Dict[str, str]) -> str:
        """Resolve image references to use consolidated images"""
        
        def replace_image_ref(match):
            alt_text = match.group(1)
            img_path = match.group(2)
            
            # Skip data URIs and URLs
            if img_path.startswith(('data:', 'http://', 'https://')):
                return match.group(0)
            
            # Resolve relative path
            if not img_path.startswith('/'):
                resolved_path = (source_path.parent / img_path).resolve()
            else:
                resolved_path = Path(img_path)
            
            # Check if we have a consolidated version
            consolidated_path = consolidated_images.get(str(resolved_path))
            if consolidated_path:
                # Use relative path from export directory
                rel_path = os.path.relpath(consolidated_path, source_path.parent)
                return f"![{alt_text}]({rel_path})"
            
            return match.group(0)
        
        return re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', replace_image_ref, content)
    
    def _resolve_internal_references(self, content: str, source_path: Path) -> str:
        """Resolve internal document references and create anchors"""
        
        # Add anchors to headers
        def add_header_anchor(match):
            level = len(match.group(1))
            title = match.group(2).strip()
            anchor = re.sub(r'[^\w\-]', '-', title.lower()).strip('-')
            return f"{'#' * level} {title} {{#{anchor}}}"
        
        content = re.sub(r'^(#{1,6})\s+(.+)$', add_header_anchor, content, flags=re.MULTILINE)
        
        # Resolve cross-references
        def resolve_cross_ref(match):
            ref_text = match.group(1)
            ref_target = match.group(2)
            
            # Convert to anchor link
            anchor = re.sub(r'[^\w\-]', '-', ref_target.lower()).strip('-')
            return f"[{ref_text}](#{anchor})"
        
        content = re.sub(r'\[([^\]]+)\]\(#([^)]+)\)', resolve_cross_ref, content)
        
        return content
    
    def _resolve_file_references(self, content: str, source_path: Path) -> str:
        """Resolve references to other files"""
        
        def resolve_file_ref(match):
            link_text = match.group(1)
            file_path = match.group(2)
            
            # Skip URLs
            if file_path.startswith(('http://', 'https://', 'mailto:')):
                return match.group(0)
            
            # Resolve relative path
            if not file_path.startswith('/'):
                resolved_path = (source_path.parent / file_path).resolve()
            else:
                resolved_path = Path(file_path)
            
            # Check if file exists
            if resolved_path.exists():
                rel_path = os.path.relpath(resolved_path, source_path.parent)
                return f"[{link_text}]({rel_path})"
            
            return match.group(0)
        
        return re.sub(r'\[([^\]]+)\]\(([^)#]+)\)', resolve_file_ref, content)

class ExportTemplate:
    """Base class for export templates"""
    
    def __init__(self, name: str, base_dir: Path):
        self.name = name
        self.base_dir = base_dir
        self.template_dir = base_dir / "export" / "templates" / name
        self.template_dir.mkdir(parents=True, exist_ok=True)
        
    def get_css_styles(self) -> str:
        """Get CSS styles for this template"""
        raise NotImplementedError
        
    def get_html_template(self) -> str:
        """Get HTML template"""
        raise NotImplementedError
        
    def process_content(self, content: str, metadata: Dict[str, Any]) -> str:
        """Process content for this template"""
        return content

class AcademicTemplate(ExportTemplate):
    """Academic template with professional styling"""
    
    def get_css_styles(self) -> str:
        return """
        @page {
            size: A4;
            margin: 2cm 2.5cm 2cm 2.5cm;
            @top-center {
                content: string(document-title);
                font-size: 10pt;
                color: #666;
            }
            @bottom-center {
                content: "Page " counter(page) " of " counter(pages);
                font-size: 10pt;
                color: #666;
            }
        }
        
        body {
            font-family: 'Times New Roman', serif;
            font-size: 11pt;
            line-height: 1.6;
            color: #333;
            margin: 0;
            padding: 0;
        }
        
        h1 {
            string-set: document-title content();
            color: #1a1a1a;
            font-size: 18pt;
            font-weight: bold;
            margin: 0 0 20pt 0;
            padding-bottom: 10pt;
            border-bottom: 2px solid #333;
            page-break-after: avoid;
        }
        
        h2 {
            color: #2c3e50;
            font-size: 14pt;
            font-weight: bold;
            margin: 24pt 0 12pt 0;
            page-break-after: avoid;
            border-bottom: 1px solid #bdc3c7;
            padding-bottom: 4pt;
        }
        
        h3 {
            color: #34495e;
            font-size: 12pt;
            font-weight: bold;
            margin: 18pt 0 10pt 0;
            page-break-after: avoid;
        }
        
        h4, h5, h6 {
            color: #7f8c8d;
            font-size: 11pt;
            font-weight: bold;
            margin: 14pt 0 8pt 0;
            page-break-after: avoid;
        }
        
        p {
            margin: 0 0 12pt 0;
            text-align: justify;
            orphans: 2;
            widows: 2;
        }
        
        img {
            max-width: 70%;
            height: auto;
            display: block;
            margin: 16pt auto;
            page-break-inside: avoid;
            border: 1px solid #e0e0e0;
            border-radius: 4px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        
        .figure-caption {
            font-size: 9pt;
            color: #666;
            font-style: italic;
            text-align: center;
            margin: 4pt 0 16pt 0;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 16pt 0;
            font-size: 10pt;
            page-break-inside: avoid;
        }
        
        th, td {
            border: 1px solid #bdc3c7;
            padding: 8pt;
            text-align: left;
            vertical-align: top;
        }
        
        th {
            background-color: #ecf0f1;
            font-weight: bold;
            color: #2c3e50;
        }
        
        tr:nth-child(even) {
            background-color: #f8f9fa;
        }
        
        ul, ol {
            margin: 12pt 0;
            padding-left: 24pt;
        }
        
        li {
            margin: 4pt 0;
            line-height: 1.6;
        }
        
        code {
            background-color: #f4f4f4;
            padding: 2pt 4pt;
            border-radius: 2px;
            font-family: 'Courier New', monospace;
            font-size: 9pt;
            color: #d63384;
        }
        
        pre {
            background-color: #f8f9fa;
            padding: 12pt;
            border-radius: 4px;
            border: 1px solid #e9ecef;
            overflow-x: auto;
            font-family: 'Courier New', monospace;
            font-size: 9pt;
            line-height: 1.4;
            page-break-inside: avoid;
            margin: 16pt 0;
        }
        
        blockquote {
            margin: 16pt 0;
            padding: 12pt 16pt;
            border-left: 4px solid #3498db;
            background-color: #f8f9fa;
            font-style: italic;
            page-break-inside: avoid;
        }
        
        .text-diagram-container {
            background-color: #f8f9fa;
            border: 2px solid #3498db;
            border-radius: 8px;
            padding: 16pt;
            margin: 16pt auto;
            max-width: 80%;
            page-break-inside: avoid;
            font-family: 'Courier New', monospace;
            font-size: 9pt;
            line-height: 1.4;
        }
        
        .diagram-title {
            font-size: 11pt;
            font-weight: bold;
            color: #2c3e50;
            margin-bottom: 8pt;
            text-align: center;
            font-family: 'Times New Roman', serif;
        }
        
        .abstract {
            background-color: #f8f9fa;
            border: 1px solid #dee2e6;
            padding: 16pt;
            margin: 20pt 0;
            font-style: italic;
            page-break-inside: avoid;
        }
        
        .keywords {
            margin: 12pt 0;
            font-weight: bold;
        }
        
        .references {
            page-break-before: auto;
        }
        
        .references h2 {
            margin-top: 40pt;
        }
        
        .reference-item {
            margin-bottom: 8pt;
            text-indent: -24pt;
            padding-left: 24pt;
        }
        
        .toc {
            page-break-after: always;
        }
        
        .toc-item {
            margin: 4pt 0;
            display: flex;
            justify-content: space-between;
        }
        
        .toc-title {
            flex-grow: 1;
        }
        
        .toc-page {
            font-weight: bold;
        }
        """
    
    def get_html_template(self) -> str:
        return """
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{title}</title>
            <style>{css}</style>
        </head>
        <body>
            <div class="document">
                {content}
            </div>
        </body>
        </html>
        """
    
    def process_content(self, content: str, metadata: Dict[str, Any]) -> str:
        """Process content for academic template"""
        
        # Add metadata section if available
        if metadata:
            meta_section = self._create_metadata_section(metadata)
            content = meta_section + "\n\n" + content
            
        # Process special sections
        content = self._process_abstract(content)
        content = self._process_keywords(content)
        content = self._process_references(content)
        
        return content
    
    def _create_metadata_section(self, metadata: Dict[str, Any]) -> str:
        """Create metadata section"""
        parts = []
        
        if metadata.get('title'):
            parts.append(f"# {metadata['title']}")
        
        meta_items = []
        for key in ['author', 'institution', 'date', 'subject']:
            if metadata.get(key):
                meta_items.append(f"**{key.title()}:** {metadata[key]}")
        
        if meta_items:
            parts.append("\n".join(meta_items))
        
        return "\n\n".join(parts)
    
    def _process_abstract(self, content: str) -> str:
        """Process abstract sections"""
        content = re.sub(
            r'^## Abstract\s*\n(.*?)(?=\n## |\n# |\Z)',
            r'<div class="abstract">\n<h2>Abstract</h2>\n\1\n</div>',
            content,
            flags=re.MULTILINE | re.DOTALL
        )
        return content
    
    def _process_keywords(self, content: str) -> str:
        """Process keywords sections"""
        content = re.sub(
            r'^\*\*Keywords:\*\*(.+)$',
            r'<div class="keywords">**Keywords:**\1</div>',
            content,
            flags=re.MULTILINE
        )
        return content
    
    def _process_references(self, content: str) -> str:
        """Process references section"""
        # Convert reference lists to proper format
        content = re.sub(
            r'^(\d+\.\s+.+)$',
            r'<div class="reference-item">\1</div>',
            content,
            flags=re.MULTILINE
        )
        return content

class PDFExporter:
    """Exports content to PDF format"""
    
    def __init__(self, template: ExportTemplate):
        self.template = template
        
    def export(self, content: str, output_path: Path, 
              metadata: Dict[str, Any] = None) -> ExportResult:
        """Export content to PDF"""
        
        start_time = time.time()
        errors = []
        warnings = []
        
        try:
            # Process content with template
            processed_content = self.template.process_content(content, metadata or {})
            
            # Convert markdown to HTML
            html_content = self._markdown_to_html(processed_content)
            
            # Create full HTML document
            full_html = self.template.get_html_template().format(
                title=metadata.get('title', 'Document') if metadata else 'Document',
                css=self.template.get_css_styles(),
                content=html_content
            )
            
            # Generate PDF using WeasyPrint
            try:
                from weasyprint import HTML, CSS
                from weasyprint.text.fonts import FontConfiguration
                
                font_config = FontConfiguration()
                css = CSS(string=self.template.get_css_styles(), font_config=font_config)
                
                HTML(string=full_html).write_pdf(
                    output_path,
                    stylesheets=[css],
                    font_config=font_config,
                    optimize_size=('fonts', 'images')
                )
                
            except ImportError:
                errors.append("WeasyPrint not available, PDF export failed")
                return ExportResult(
                    success=False,
                    output_files=[],
                    format_type='pdf',
                    file_size_mb=0,
                    processing_time=time.time() - start_time,
                    validation_score=0,
                    errors=errors,
                    warnings=warnings,
                    metadata={}
                )
            
            # Validate output
            if output_path.exists():
                file_size = output_path.stat().st_size / (1024 * 1024)
                validation_score = self._validate_pdf(output_path)
                
                return ExportResult(
                    success=True,
                    output_files=[str(output_path)],
                    format_type='pdf',
                    file_size_mb=file_size,
                    processing_time=time.time() - start_time,
                    validation_score=validation_score,
                    errors=errors,
                    warnings=warnings,
                    metadata={'pages': self._count_pdf_pages(output_path)}
                )
            else:
                errors.append("PDF file was not created")
                
        except Exception as e:
            errors.append(f"PDF export failed: {str(e)}")
        
        return ExportResult(
            success=False,
            output_files=[],
            format_type='pdf',
            file_size_mb=0,
            processing_time=time.time() - start_time,
            validation_score=0,
            errors=errors,
            warnings=warnings,
            metadata={}
        )
    
    def _markdown_to_html(self, content: str) -> str:
        """Convert markdown to HTML"""
        try:
            import markdown
            md = markdown.Markdown(extensions=['tables', 'fenced_code', 'toc'])
            return md.convert(content)
        except ImportError:
            # Simple fallback conversion
            return self._simple_markdown_to_html(content)
    
    def _simple_markdown_to_html(self, content: str) -> str:
        """Simple markdown to HTML conversion"""
        lines = content.split('\n')
        html_lines = []
        in_code_block = False
        
        for line in lines:
            if line.startswith('```'):
                if in_code_block:
                    html_lines.append('</pre>')
                    in_code_block = False
                else:
                    html_lines.append('<pre>')
                    in_code_block = True
            elif in_code_block:
                html_lines.append(line)
            elif line.startswith('# '):
                html_lines.append(f'<h1>{line[2:]}</h1>')
            elif line.startswith('## '):
                html_lines.append(f'<h2>{line[3:]}</h2>')
            elif line.startswith('### '):
                html_lines.append(f'<h3>{line[4:]}</h3>')
            elif line.strip() == '':
                html_lines.append('<br>')
            else:
                html_lines.append(f'<p>{line}</p>')
        
        return '\n'.join(html_lines)
    
    def _validate_pdf(self, pdf_path: Path) -> float:
        """Validate PDF quality"""
        try:
            # Basic validation: file exists and has reasonable size
            if not pdf_path.exists():
                return 0.0
            
            file_size = pdf_path.stat().st_size
            if file_size < 1024:  # Less than 1KB
                return 0.1
            
            # Try to count pages
            page_count = self._count_pdf_pages(pdf_path)
            if page_count == 0:
                return 0.2
            
            # Basic score based on file size and pages
            score = min(0.8 + (page_count * 0.05), 1.0)
            return score
            
        except Exception:
            return 0.3
    
    def _count_pdf_pages(self, pdf_path: Path) -> int:
        """Count pages in PDF"""
        try:
            # Try using PyPDF2
            import PyPDF2
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return len(reader.pages)
        except ImportError:
            pass
        
        try:
            # Try using pdfplumber
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                return len(pdf.pages)
        except ImportError:
            pass
        
        # Fallback: estimate based on file size
        file_size = pdf_path.stat().st_size
        estimated_pages = max(1, file_size // (50 * 1024))  # ~50KB per page
        return estimated_pages

class HTMLExporter:
    """Exports content to HTML format"""
    
    def __init__(self, template: ExportTemplate):
        self.template = template
        
    def export(self, content: str, output_path: Path,
              metadata: Dict[str, Any] = None) -> ExportResult:
        """Export content to HTML"""
        
        start_time = time.time()
        errors = []
        warnings = []
        
        try:
            # Process content with template
            processed_content = self.template.process_content(content, metadata or {})
            
            # Convert markdown to HTML
            html_content = self._markdown_to_html(processed_content)
            
            # Create full HTML document
            full_html = self.template.get_html_template().format(
                title=metadata.get('title', 'Document') if metadata else 'Document',
                css=self.template.get_css_styles(),
                content=html_content
            )
            
            # Write HTML file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
            
            # Validate output
            if output_path.exists():
                file_size = output_path.stat().st_size / (1024 * 1024)
                validation_score = self._validate_html(output_path)
                
                return ExportResult(
                    success=True,
                    output_files=[str(output_path)],
                    format_type='html',
                    file_size_mb=file_size,
                    processing_time=time.time() - start_time,
                    validation_score=validation_score,
                    errors=errors,
                    warnings=warnings,
                    metadata={'word_count': len(content.split())}
                )
            else:
                errors.append("HTML file was not created")
                
        except Exception as e:
            errors.append(f"HTML export failed: {str(e)}")
        
        return ExportResult(
            success=False,
            output_files=[],
            format_type='html',
            file_size_mb=0,
            processing_time=time.time() - start_time,
            validation_score=0,
            errors=errors,
            warnings=warnings,
            metadata={}
        )
    
    def _markdown_to_html(self, content: str) -> str:
        """Convert markdown to HTML with enhanced features"""
        try:
            import markdown
            md = markdown.Markdown(extensions=[
                'tables', 
                'fenced_code', 
                'toc',
                'codehilite',
                'attr_list'
            ])
            return md.convert(content)
        except ImportError:
            return self._simple_markdown_to_html(content)
    
    def _simple_markdown_to_html(self, content: str) -> str:
        """Simple markdown to HTML conversion"""
        # Enhanced version of the simple converter
        lines = content.split('\n')
        html_lines = []
        in_code_block = False
        in_list = False
        
        for line in lines:
            if line.startswith('```'):
                if in_code_block:
                    html_lines.append('</pre>')
                    in_code_block = False
                else:
                    html_lines.append('<pre><code>')
                    in_code_block = True
            elif in_code_block:
                html_lines.append(line)
            elif line.startswith('# '):
                html_lines.append(f'<h1>{line[2:]}</h1>')
            elif line.startswith('## '):
                html_lines.append(f'<h2>{line[3:]}</h2>')
            elif line.startswith('### '):
                html_lines.append(f'<h3>{line[4:]}</h3>')
            elif line.startswith('- ') or line.startswith('* '):
                if not in_list:
                    html_lines.append('<ul>')
                    in_list = True
                html_lines.append(f'<li>{line[2:]}</li>')
            elif line.strip() == '':
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append('')
            else:
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                # Process inline formatting
                processed_line = self._process_inline_formatting(line)
                html_lines.append(f'<p>{processed_line}</p>')
        
        if in_code_block:
            html_lines.append('</code></pre>')
        if in_list:
            html_lines.append('</ul>')
        
        return '\n'.join(html_lines)
    
    def _process_inline_formatting(self, text: str) -> str:
        """Process inline markdown formatting"""
        # Bold
        text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'__(.*?)__', r'<strong>\1</strong>', text)
        
        # Italic
        text = re.sub(r'\*(.*?)\*', r'<em>\1</em>', text)
        text = re.sub(r'_(.*?)_', r'<em>\1</em>', text)
        
        # Code
        text = re.sub(r'`(.*?)`', r'<code>\1</code>', text)
        
        # Links
        text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', text)
        
        # Images
        text = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', r'<img src="\2" alt="\1" />', text)
        
        return text
    
    def _validate_html(self, html_path: Path) -> float:
        """Validate HTML quality"""
        try:
            if not html_path.exists():
                return 0.0
            
            with open(html_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            score = 0.5  # Base score
            
            # Check for basic HTML structure
            if '<html>' in content and '</html>' in content:
                score += 0.2
            if '<head>' in content and '</head>' in content:
                score += 0.1
            if '<body>' in content and '</body>' in content:
                score += 0.1
            if '<title>' in content:
                score += 0.1
            
            return min(score, 1.0)
            
        except Exception:
            return 0.3

class DOCXExporter:
    """Exports content to DOCX format"""
    
    def __init__(self, template: ExportTemplate):
        self.template = template
        
    def export(self, content: str, output_path: Path,
              metadata: Dict[str, Any] = None) -> ExportResult:
        """Export content to DOCX"""
        
        start_time = time.time()
        errors = []
        warnings = []
        
        try:
            # Try using python-docx
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Add metadata
            if metadata:
                self._add_metadata_to_doc(doc, metadata)
            
            # Process content
            processed_content = self.template.process_content(content, metadata or {})
            
            # Parse markdown and add to document
            self._markdown_to_docx(processed_content, doc)
            
            # Save document
            doc.save(output_path)
            
            if output_path.exists():
                file_size = output_path.stat().st_size / (1024 * 1024)
                validation_score = self._validate_docx(output_path)
                
                return ExportResult(
                    success=True,
                    output_files=[str(output_path)],
                    format_type='docx',
                    file_size_mb=file_size,
                    processing_time=time.time() - start_time,
                    validation_score=validation_score,
                    errors=errors,
                    warnings=warnings,
                    metadata={'paragraphs': len(doc.paragraphs)}
                )
            else:
                errors.append("DOCX file was not created")
                
        except ImportError:
            errors.append("python-docx not available, DOCX export failed")
        except Exception as e:
            errors.append(f"DOCX export failed: {str(e)}")
        
        return ExportResult(
            success=False,
            output_files=[],
            format_type='docx',
            file_size_mb=0,
            processing_time=time.time() - start_time,
            validation_score=0,
            errors=errors,
            warnings=warnings,
            metadata={}
        )
    
    def _add_metadata_to_doc(self, doc, metadata: Dict[str, Any]) -> None:
        """Add metadata to document properties"""
        props = doc.core_properties
        
        if metadata.get('title'):
            props.title = metadata['title']
        if metadata.get('author'):
            props.author = metadata['author']
        if metadata.get('subject'):
            props.subject = metadata['subject']
        if metadata.get('keywords'):
            props.keywords = metadata['keywords']
    
    def _markdown_to_docx(self, content: str, doc) -> None:
        """Convert markdown content to DOCX document"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            if line.startswith('# '):
                # Heading 1
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                # Heading 2
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # Heading 3
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet point
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                # Regular paragraph
                doc.add_paragraph(line)
    
    def _validate_docx(self, docx_path: Path) -> float:
        """Validate DOCX quality"""
        try:
            if not docx_path.exists():
                return 0.0
            
            # Basic validation
            file_size = docx_path.stat().st_size
            if file_size < 1024:  # Less than 1KB
                return 0.1
            
            # Try to open and check basic structure
            try:
                from docx import Document
                doc = Document(docx_path)
                
                if len(doc.paragraphs) > 0:
                    return 0.8
                else:
                    return 0.3
            except ImportError:
                return 0.5
            
        except Exception:
            return 0.2

class ExportQualityValidator:
    """Validates export quality and provides optimization suggestions"""
    
    def __init__(self):
        self.validation_criteria = {
            'file_size': {'min': 0.001, 'max': 50, 'weight': 0.2},  # MB
            'content_structure': {'weight': 0.3},
            'image_quality': {'weight': 0.2},
            'reference_integrity': {'weight': 0.2},
            'format_compliance': {'weight': 0.1}
        }
    
    def validate_export(self, export_result: ExportResult, 
                       original_content: str) -> Dict[str, Any]:
        """Validate export quality and provide score"""
        
        validation_results = {
            'overall_score': 0.0,
            'criteria_scores': {},
            'issues': [],
            'suggestions': [],
            'passed': False
        }
        
        try:
            # Validate file size
            size_score = self._validate_file_size(export_result.file_size_mb)
            validation_results['criteria_scores']['file_size'] = size_score
            
            # Validate content structure
            structure_score = self._validate_content_structure(
                export_result, original_content
            )
            validation_results['criteria_scores']['content_structure'] = structure_score
            
            # Validate format compliance
            format_score = self._validate_format_compliance(export_result)
            validation_results['criteria_scores']['format_compliance'] = format_score
            
            # Calculate overall score
            overall_score = 0.0
            for criterion, score in validation_results['criteria_scores'].items():
                weight = self.validation_criteria[criterion]['weight']
                overall_score += score * weight
            
            validation_results['overall_score'] = overall_score
            validation_results['passed'] = overall_score >= 0.7
            
            # Generate suggestions
            validation_results['suggestions'] = self._generate_suggestions(
                validation_results['criteria_scores']
            )
            
        except Exception as e:
            validation_results['issues'].append(f"Validation error: {str(e)}")
        
        return validation_results
    
    def _validate_file_size(self, file_size_mb: float) -> float:
        """Validate file size"""
        criteria = self.validation_criteria['file_size']
        
        if file_size_mb < criteria['min']:
            return 0.1  # Too small, likely empty or corrupted
        elif file_size_mb > criteria['max']:
            return 0.5  # Too large, but not necessarily bad
        else:
            # Good size range
            return 0.9
    
    def _validate_content_structure(self, export_result: ExportResult,
                                   original_content: str) -> float:
        """Validate content structure preservation"""
        
        # Basic checks
        if not export_result.success:
            return 0.0
        
        if not export_result.output_files:
            return 0.0
        
        # Count structural elements in original
        original_headers = len(re.findall(r'^#+\s', original_content, re.MULTILINE))
        original_images = len(re.findall(r'!\[.*?\]\(.*?\)', original_content))
        
        # Basic score based on successful export
        score = 0.6
        
        # Bonus for preserving structure (would need to parse output file)
        if original_headers > 0:
            score += 0.2
        if original_images > 0:
            score += 0.2
        
        return min(score, 1.0)
    
    def _validate_format_compliance(self, export_result: ExportResult) -> float:
        """Validate format-specific compliance"""
        
        if not export_result.success:
            return 0.0
        
        format_type = export_result.format_type.lower()
        
        if format_type == 'pdf':
            # Check if PDF has reasonable metadata
            if export_result.metadata.get('pages', 0) > 0:
                return 0.9
            else:
                return 0.5
        elif format_type == 'html':
            # Check word count preservation
            if export_result.metadata.get('word_count', 0) > 0:
                return 0.8
            else:
                return 0.5
        elif format_type == 'docx':
            # Check paragraph count
            if export_result.metadata.get('paragraphs', 0) > 0:
                return 0.8
            else:
                return 0.5
        
        return 0.7  # Default for unknown formats
    
    def _generate_suggestions(self, criteria_scores: Dict[str, float]) -> List[str]:
        """Generate optimization suggestions"""
        suggestions = []
        
        for criterion, score in criteria_scores.items():
            if score < 0.7:
                if criterion == 'file_size':
                    suggestions.append(
                        "Consider optimizing images or reducing content length"
                    )
                elif criterion == 'content_structure':
                    suggestions.append(
                        "Review content processing to ensure structure preservation"
                    )
                elif criterion == 'format_compliance':
                    suggestions.append(
                        f"Check format-specific requirements for {criterion}"
                    )
        
        return suggestions

class ExportSystemTool(Tool):
    """Main export system tool for the academic agent"""
    
    name = "export_system_tool"
    description = "Comprehensive export system supporting PDF, HTML, DOCX with image consolidation and quality validation"
    inputs = {
        "content_paths": {
            "type": "array",
            "description": "Paths to content files or directories to export",
        },
        "output_directory": {
            "type": "string",
            "description": "Directory for export outputs",
        },
        "export_config": {
            "type": "object",
            "description": "Export configuration (format, template, options)",
            "nullable": True,
        },
        "batch_mode": {
            "type": "boolean",
            "description": "Enable batch processing mode",
            "nullable": True,
        },
    }
    outputs = {
        "export_results": {
            "type": "array",
            "description": "Results for each exported file",
        },
        "consolidation_result": {
            "type": "object",
            "description": "Image consolidation results",
        },
        "validation_results": {
            "type": "array",
            "description": "Quality validation results",
        },
        "summary": {
            "type": "object",
            "description": "Export operation summary",
        },
    }
    output_type = "object"
    
    def __init__(self, base_dir: Optional[Path] = None):
        super().__init__()
        self.base_dir = base_dir or Path.cwd()
        self.export_dir = self.base_dir / "export"
        self.export_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize components
        self.image_consolidator = ImageConsolidator(self.base_dir)
        self.reference_resolver = ReferenceResolver(self.base_dir)
        self.quality_validator = ExportQualityValidator()
        
        # Available templates
        self.templates = {
            'academic': AcademicTemplate('academic', self.base_dir)
        }
        
        # Available exporters
        self.exporters = {}
        
    def forward(self, content_paths: List[str], output_directory: str,
                export_config: Optional[Dict[str, Any]] = None,
                batch_mode: bool = False) -> Dict[str, Any]:
        """Main export function"""
        
        start_time = time.time()
        
        # Parse configuration
        config = ExportConfig(**(export_config or {}))
        
        # Ensure output directory exists
        output_dir = Path(output_directory)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Convert paths to Path objects
        content_path_objects = [Path(p) for p in content_paths]
        
        try:
            # Step 1: Consolidate images
            consolidation_result = self.image_consolidator.consolidate_images(
                content_path_objects, config.image_sizing
            )
            
            # Step 2: Process each content file
            export_results = []
            
            if batch_mode:
                export_results = self._batch_export(
                    content_path_objects, output_dir, config, consolidation_result
                )
            else:
                for content_path in content_path_objects:
                    if content_path.exists():
                        result = self._export_single_file(
                            content_path, output_dir, config, consolidation_result
                        )
                        export_results.append(result)
            
            # Step 3: Validate results
            validation_results = []
            for result in export_results:
                if result.success:
                    # Read original content for validation
                    try:
                        with open(content_path_objects[0], 'r', encoding='utf-8') as f:
                            original_content = f.read()
                    except:
                        original_content = ""
                    
                    validation = self.quality_validator.validate_export(
                        result, original_content
                    )
                    validation_results.append(validation)
            
            # Step 4: Generate summary
            summary = self._generate_summary(
                export_results, consolidation_result, validation_results, start_time
            )
            
            return {
                "export_results": [asdict(r) for r in export_results],
                "consolidation_result": asdict(consolidation_result),
                "validation_results": validation_results,
                "summary": summary
            }
            
        except Exception as e:
            return {
                "export_results": [],
                "consolidation_result": asdict(ImageConsolidationResult(
                    consolidated_images={}, total_size_mb=0, 
                    optimization_ratio=0, errors=[str(e)]
                )),
                "validation_results": [],
                "summary": {
                    "success": False,
                    "error": str(e),
                    "processing_time": time.time() - start_time
                }
            }
    
    def _export_single_file(self, content_path: Path, output_dir: Path,
                           config: ExportConfig, 
                           consolidation_result: ImageConsolidationResult) -> ExportResult:
        """Export a single file"""
        
        try:
            # Read content
            with open(content_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Resolve references
            if config.resolve_references:
                content = self.reference_resolver.resolve_references(
                    content, content_path, consolidation_result.consolidated_images
                )
            
            # Get template
            template = self.templates.get(config.template_name)
            if not template:
                template = self.templates['academic']  # Default
            
            # Determine formats to export
            formats = []
            if config.output_format == 'all':
                formats = ['pdf', 'html', 'docx']
            else:
                formats = [config.output_format]
            
            # Export to each format
            all_results = []
            for fmt in formats:
                exporter = self._get_exporter(fmt, template)
                if exporter:
                    output_path = output_dir / f"{content_path.stem}.{fmt}"
                    
                    # Prepare metadata
                    metadata = config.metadata or {}
                    metadata.update({
                        'title': content_path.stem.replace('_', ' ').title(),
                        'source_file': str(content_path),
                        'export_time': datetime.now().isoformat()
                    })
                    
                    result = exporter.export(content, output_path, metadata)
                    all_results.append(result)
            
            # Return the primary result (first format)
            return all_results[0] if all_results else ExportResult(
                success=False, output_files=[], format_type=config.output_format,
                file_size_mb=0, processing_time=0, validation_score=0,
                errors=["No compatible exporter found"], warnings=[], metadata={}
            )
            
        except Exception as e:
            return ExportResult(
                success=False, output_files=[], format_type=config.output_format,
                file_size_mb=0, processing_time=0, validation_score=0,
                errors=[f"Export failed: {str(e)}"], warnings=[], metadata={}
            )
    
    def _batch_export(self, content_paths: List[Path], output_dir: Path,
                     config: ExportConfig, 
                     consolidation_result: ImageConsolidationResult) -> List[ExportResult]:
        """Export multiple files in batch mode"""
        
        results = []
        
        # Process files concurrently if possible
        for content_path in content_paths:
            if content_path.exists():
                result = self._export_single_file(
                    content_path, output_dir, config, consolidation_result
                )
                results.append(result)
        
        return results
    
    def _get_exporter(self, format_type: str, template: ExportTemplate):
        """Get appropriate exporter for format"""
        
        if format_type not in self.exporters:
            if format_type == 'pdf':
                self.exporters[format_type] = PDFExporter(template)
            elif format_type == 'html':
                self.exporters[format_type] = HTMLExporter(template)
            elif format_type == 'docx':
                self.exporters[format_type] = DOCXExporter(template)
            else:
                return None
        
        return self.exporters[format_type]
    
    def _generate_summary(self, export_results: List[ExportResult],
                         consolidation_result: ImageConsolidationResult,
                         validation_results: List[Dict[str, Any]],
                         start_time: float) -> Dict[str, Any]:
        """Generate export summary"""
        
        successful_exports = sum(1 for r in export_results if r.success)
        total_size_mb = sum(r.file_size_mb for r in export_results if r.success)
        avg_validation_score = (
            sum(v['overall_score'] for v in validation_results) / len(validation_results)
            if validation_results else 0
        )
        
        return {
            "success": successful_exports > 0,
            "total_files_processed": len(export_results),
            "successful_exports": successful_exports,
            "failed_exports": len(export_results) - successful_exports,
            "total_output_size_mb": total_size_mb,
            "image_consolidation": {
                "images_processed": len(consolidation_result.consolidated_images),
                "total_image_size_mb": consolidation_result.total_size_mb,
                "optimization_ratio": consolidation_result.optimization_ratio
            },
            "average_validation_score": avg_validation_score,
            "processing_time_seconds": time.time() - start_time,
            "output_formats": list(set(r.format_type for r in export_results)),
            "all_output_files": [f for r in export_results for f in r.output_files]
        }

# Integration function
def setup_export_system(base_dir: Optional[Path] = None) -> ExportSystemTool:
    """Set up the export system tool"""
    return ExportSystemTool(base_dir)

if __name__ == "__main__":
    # Example usage
    export_system = ExportSystemTool()
    
    # Test configuration
    config = {
        "output_format": "pdf",
        "template_name": "academic",
        "image_sizing": "medium",
        "include_diagrams": True,
        "consolidate_images": True,
        "resolve_references": True
    }
    
    # Example export
    result = export_system.forward(
        content_paths=["sample_content.md"],
        output_directory="./export_test",
        export_config=config,
        batch_mode=False
    )
    
    print("Export System Results:")
    print(f"Success: {result['summary']['success']}")
    print(f"Files processed: {result['summary']['total_files_processed']}")
    print(f"Total output size: {result['summary']['total_output_size_mb']:.2f} MB")