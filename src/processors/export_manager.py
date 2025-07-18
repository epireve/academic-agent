#!/usr/bin/env python3
"""
Enhanced Export Manager for Academic Agent

This module provides comprehensive export functionality with support for multiple formats,
consolidated image handling, local reference resolution, and quality validation.

Key Features:
- Multiple export formats (PDF, HTML, DOCX, EPUB)
- Consolidated image handling with optimization
- Local reference resolution and linking
- Template-based export with customizable styling
- Quality validation integration
- Batch export capabilities
- Content Management System integration
"""

import os
import json
import re
import time
import asyncio
import subprocess
import tempfile
import shutil
import uuid
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple, Union
from dataclasses import dataclass, asdict
from datetime import datetime
from abc import ABC, abstractmethod
import hashlib
import base64
from urllib.parse import quote
import zipfile
import mimetypes

# PDF generation libraries
try:
    from weasyprint import HTML, CSS
    from weasyprint.text.fonts import FontConfiguration
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

# Alternative PDF library
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# DOCX library
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# EPUB library
try:
    from ebooklib import epub
    EBOOKLIB_AVAILABLE = True
except ImportError:
    EBOOKLIB_AVAILABLE = False

# Markdown library
try:
    import markdown
    from markdown.extensions import Extension
    from markdown.extensions.toc import TocExtension
    from markdown.extensions.tables import TableExtension
    from markdown.extensions.fenced_code import FencedCodeExtension
    from markdown.extensions.codehilite import CodeHiliteExtension
    from markdown.extensions.attr_list import AttrListExtension
    from markdown.extensions.def_list import DefListExtension
    from markdown.extensions.footnotes import FootnoteExtension
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

# Image processing
try:
    from PIL import Image as PILImage
    from PIL import ImageDraw, ImageFont
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False


@dataclass
class ExportConfig:
    """Configuration for export operations"""
    output_format: str  # 'pdf', 'html', 'docx', 'epub', 'all'
    template_name: str = 'academic'
    image_sizing: str = 'medium'  # 'small', 'medium', 'large', 'original'
    include_diagrams: bool = True
    diagram_format: str = 'png'  # 'png', 'svg', 'text'
    quality_level: str = 'high'  # 'low', 'medium', 'high'
    consolidate_images: bool = True
    embed_images: bool = True
    resolve_references: bool = True
    optimize_for_print: bool = True
    include_toc: bool = True
    include_index: bool = True
    include_metadata: bool = True
    custom_css: Optional[str] = None
    custom_fonts: Optional[List[str]] = None
    metadata: Dict[str, Any] = None
    page_size: str = 'A4'  # 'A4', 'Letter', 'Legal'
    margin_size: str = 'normal'  # 'narrow', 'normal', 'wide'
    enable_compression: bool = True
    watermark_text: Optional[str] = None
    header_text: Optional[str] = None
    footer_text: Optional[str] = None


@dataclass
class ExportResult:
    """Result of an export operation"""
    success: bool
    output_files: List[str]
    format_type: str
    file_size_mb: float
    processing_time: float
    validation_score: float
    errors: List[str]
    warnings: List[str]
    metadata: Dict[str, Any]
    quality_report: Optional[Dict[str, Any]] = None
    optimization_report: Optional[Dict[str, Any]] = None


@dataclass
class ImageConsolidationResult:
    """Result of image consolidation"""
    consolidated_images: Dict[str, str]  # original_path -> consolidated_path
    total_size_mb: float
    optimization_ratio: float
    processed_count: int
    skipped_count: int
    errors: List[str]
    image_metadata: Dict[str, Dict[str, Any]] = None


class ImageConsolidator:
    """Enhanced image consolidation and optimization"""
    
    def __init__(self, base_dir: Path):
        self.base_dir = base_dir
        self.consolidated_dir = base_dir / "export" / "images"
        self.consolidated_dir.mkdir(parents=True, exist_ok=True)
        self.image_cache = {}
        
    def consolidate_images(self, content_paths: List[Path], 
                          sizing: str = 'medium',
                          quality_config: Dict[str, Any] = None) -> ImageConsolidationResult:
        """Consolidate and optimize images from multiple content sources"""
        
        consolidated_images = {}
        total_original_size = 0
        total_consolidated_size = 0
        processed_count = 0
        skipped_count = 0
        errors = []
        image_metadata = {}
        
        # Enhanced size configurations with quality settings
        size_configs = {
            'small': {'max_width': 400, 'quality': 70, 'dpi': 72},
            'medium': {'max_width': 600, 'quality': 80, 'dpi': 96},
            'large': {'max_width': 800, 'quality': 85, 'dpi': 150},
            'original': {'max_width': None, 'quality': 90, 'dpi': 300}
        }
        
        size_config = size_configs.get(sizing, size_configs['medium'])
        if quality_config:
            size_config.update(quality_config)
        
        # Find all images in content
        image_paths = self._find_all_images(content_paths)
        
        for img_path in image_paths:
            try:
                # Check cache first
                cache_key = self._get_cache_key(img_path, size_config)
                if cache_key in self.image_cache:
                    consolidated_images[str(img_path)] = self.image_cache[cache_key]
                    skipped_count += 1
                    continue
                
                # Calculate hash for unique naming
                with open(img_path, 'rb') as f:
                    content = f.read()
                    content_hash = hashlib.md5(content).hexdigest()[:8]
                
                # Generate consolidated filename
                ext = img_path.suffix.lower()
                consolidated_name = f"{content_hash}_{img_path.stem}_{sizing}{ext}"
                consolidated_path = self.consolidated_dir / consolidated_name
                
                # Track original size
                total_original_size += img_path.stat().st_size
                
                # Optimize and copy image
                optimization_result = self._optimize_image(
                    img_path, consolidated_path, size_config
                )
                
                if optimization_result['success']:
                    consolidated_images[str(img_path)] = str(consolidated_path)
                    total_consolidated_size += consolidated_path.stat().st_size
                    processed_count += 1
                    
                    # Store metadata
                    image_metadata[str(consolidated_path)] = {
                        'original_path': str(img_path),
                        'original_size': img_path.stat().st_size,
                        'consolidated_size': consolidated_path.stat().st_size,
                        'dimensions': optimization_result.get('dimensions'),
                        'format': optimization_result.get('format'),
                        'optimization_ratio': optimization_result.get('ratio', 0),
                        'processing_time': optimization_result.get('time', 0)
                    }
                    
                    # Update cache
                    self.image_cache[cache_key] = str(consolidated_path)
                else:
                    # Fallback: copy original
                    shutil.copy2(img_path, consolidated_path)
                    consolidated_images[str(img_path)] = str(consolidated_path)
                    total_consolidated_size += consolidated_path.stat().st_size
                    processed_count += 1
                    
            except Exception as e:
                errors.append(f"Failed to consolidate {img_path}: {str(e)}")
        
        optimization_ratio = (
            (total_original_size - total_consolidated_size) / total_original_size
            if total_original_size > 0 else 0
        )
        
        return ImageConsolidationResult(
            consolidated_images=consolidated_images,
            total_size_mb=total_consolidated_size / (1024 * 1024),
            optimization_ratio=optimization_ratio,
            processed_count=processed_count,
            skipped_count=skipped_count,
            errors=errors,
            image_metadata=image_metadata
        )
    
    def _get_cache_key(self, img_path: Path, config: Dict[str, Any]) -> str:
        """Generate cache key for image"""
        config_str = json.dumps(config, sort_keys=True)
        return f"{img_path}_{config_str}"
    
    def _find_all_images(self, content_paths: List[Path]) -> List[Path]:
        """Find all image files referenced in content"""
        image_paths = set()
        image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.svg', '.bmp', '.webp', '.tiff', '.ico'}
        
        for content_path in content_paths:
            if content_path.is_file():
                # Check content for image references
                try:
                    with open(content_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        
                    # Find markdown image references
                    img_refs = re.findall(r'!\[[^\]]*\]\(([^)]+)\)', content)
                    
                    # Find HTML image references
                    html_img_refs = re.findall(r'<img[^>]+src=["\']([^"\'>]+)["\']', content)
                    img_refs.extend(html_img_refs)
                    
                    for ref in img_refs:
                        if not ref.startswith(('http://', 'https://', 'data:', 'ftp://')):
                            # Resolve relative path
                            img_path = content_path.parent / ref
                            if img_path.exists() and img_path.suffix.lower() in image_extensions:
                                image_paths.add(img_path.resolve())
                                
                except Exception:
                    pass
                    
            elif content_path.is_dir():
                # Scan directory for images
                for img_path in content_path.rglob('*'):
                    if img_path.suffix.lower() in image_extensions:
                        image_paths.add(img_path)
        
        return sorted(list(image_paths))
    
    def _optimize_image(self, source_path: Path, target_path: Path, 
                       config: Dict[str, Any]) -> Dict[str, Any]:
        """Optimize an image file with enhanced processing"""
        result = {'success': False}
        start_time = time.time()
        
        try:
            if not PILLOW_AVAILABLE:
                return result
            
            with PILImage.open(source_path) as img:
                # Store original dimensions
                original_dimensions = img.size
                result['dimensions'] = {'original': original_dimensions}
                
                # Convert to RGB if necessary
                if img.mode in ('RGBA', 'LA', 'P'):
                    # Create white background
                    background = PILImage.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                    img = background
                
                # Apply DPI settings
                if 'dpi' in config:
                    img.info['dpi'] = (config['dpi'], config['dpi'])
                
                # Resize if needed
                if config['max_width'] and img.width > config['max_width']:
                    ratio = config['max_width'] / img.width
                    new_height = int(img.height * ratio)
                    img = img.resize((config['max_width'], new_height), PILImage.Resampling.LANCZOS)
                    result['dimensions']['resized'] = (config['max_width'], new_height)
                
                # Apply additional optimizations
                if source_path.suffix.lower() in ['.png', '.gif']:
                    # Optimize PNG/GIF
                    img.save(target_path, optimize=True, quality=config['quality'])
                else:
                    # Optimize JPEG
                    img.save(target_path, 'JPEG', quality=config['quality'], optimize=True, progressive=True)
                
                result['format'] = img.format
                result['success'] = True
                result['ratio'] = 1 - (target_path.stat().st_size / source_path.stat().st_size)
                result['time'] = time.time() - start_time
                
            return result
            
        except Exception as e:
            result['error'] = str(e)
            return result


class ReferenceResolver:
    """Enhanced reference resolution with link validation"""
    
    def __init__(self, base_dir: Path):
        self.base_dir = base_dir
        self.reference_map = {}
        self.broken_links = []
        
    def resolve_references(self, content: str, source_path: Path,
                          consolidated_images: Dict[str, str],
                          validate_links: bool = True) -> Tuple[str, List[str]]:
        """Resolve all references in content and return updated content with warnings"""
        
        warnings = []
        
        # Resolve image references
        content, img_warnings = self._resolve_image_references(
            content, source_path, consolidated_images
        )
        warnings.extend(img_warnings)
        
        # Resolve internal document references
        content, int_warnings = self._resolve_internal_references(
            content, source_path
        )
        warnings.extend(int_warnings)
        
        # Resolve external file references
        content, ext_warnings = self._resolve_file_references(
            content, source_path, validate_links
        )
        warnings.extend(ext_warnings)
        
        # Resolve cross-document references
        content, cross_warnings = self._resolve_cross_references(
            content, source_path
        )
        warnings.extend(cross_warnings)
        
        return content, warnings
    
    def _resolve_image_references(self, content: str, source_path: Path,
                                 consolidated_images: Dict[str, str]) -> Tuple[str, List[str]]:
        """Resolve image references to use consolidated images"""
        warnings = []
        
        def replace_image_ref(match):
            alt_text = match.group(1)
            img_path = match.group(2)
            
            # Skip data URIs and URLs
            if img_path.startswith(('data:', 'http://', 'https://', 'ftp://')):
                return match.group(0)
            
            # Resolve relative path
            if not img_path.startswith('/'):
                resolved_path = (source_path.parent / img_path).resolve()
            else:
                resolved_path = Path(img_path)
            
            # Check if we have a consolidated version
            consolidated_path = consolidated_images.get(str(resolved_path))
            if consolidated_path:
                # Use relative path from export directory
                try:
                    rel_path = os.path.relpath(consolidated_path, source_path.parent)
                    # Convert to forward slashes for consistency
                    rel_path = rel_path.replace('\\', '/')
                    return f"![{alt_text}]({rel_path})"
                except ValueError:
                    # If relative path can't be computed, use absolute
                    return f"![{alt_text}](file://{consolidated_path})"
            else:
                # Image not found in consolidation
                if not resolved_path.exists():
                    warnings.append(f"Missing image: {img_path}")
                    self.broken_links.append({'type': 'image', 'path': img_path})
            
            return match.group(0)
        
        # Process markdown images
        content = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', replace_image_ref, content)
        
        # Process HTML images
        def replace_html_img(match):
            full_tag = match.group(0)
            src = match.group(1)
            
            if src.startswith(('data:', 'http://', 'https://', 'ftp://')):
                return full_tag
            
            if not src.startswith('/'):
                resolved_path = (source_path.parent / src).resolve()
            else:
                resolved_path = Path(src)
            
            consolidated_path = consolidated_images.get(str(resolved_path))
            if consolidated_path:
                try:
                    rel_path = os.path.relpath(consolidated_path, source_path.parent)
                    rel_path = rel_path.replace('\\', '/')
                    return full_tag.replace(src, rel_path)
                except ValueError:
                    return full_tag.replace(src, f"file://{consolidated_path}")
            elif not resolved_path.exists():
                warnings.append(f"Missing HTML image: {src}")
                self.broken_links.append({'type': 'html_image', 'path': src})
            
            return full_tag
        
        content = re.sub(r'<img[^>]+src=["\']([^"\'>]+)["\'][^>]*>', replace_html_img, content)
        
        return content, warnings
    
    def _resolve_internal_references(self, content: str, source_path: Path) -> Tuple[str, List[str]]:
        """Resolve internal document references and create anchors"""
        warnings = []
        
        # Track all headers for validation
        headers = {}
        
        # Add anchors to headers and track them
        def add_header_anchor(match):
            level = len(match.group(1))
            title = match.group(2).strip()
            anchor = re.sub(r'[^\w\-]', '-', title.lower()).strip('-')
            
            # Handle duplicate anchors
            if anchor in headers:
                counter = 1
                while f"{anchor}-{counter}" in headers:
                    counter += 1
                anchor = f"{anchor}-{counter}"
            
            headers[anchor] = title
            return f"{'#' * level} {title} {{#{anchor}}}"
        
        content = re.sub(r'^(#{1,6})\s+(.+)$', add_header_anchor, content, flags=re.MULTILINE)
        
        # Validate and resolve cross-references
        def validate_cross_ref(match):
            ref_text = match.group(1)
            ref_target = match.group(2)
            
            # Clean up the reference target
            anchor = re.sub(r'[^\w\-]', '-', ref_target.lower()).strip('-')
            
            # Check if anchor exists
            if anchor not in headers and not any(anchor.startswith(h) for h in headers):
                warnings.append(f"Broken internal reference: #{ref_target}")
                self.broken_links.append({'type': 'internal', 'anchor': ref_target})
            
            return f"[{ref_text}](#{anchor})"
        
        content = re.sub(r'\[([^\]]+)\]\(#([^)]+)\)', validate_cross_ref, content)
        
        return content, warnings
    
    def _resolve_file_references(self, content: str, source_path: Path,
                                validate_links: bool) -> Tuple[str, List[str]]:
        """Resolve references to other files"""
        warnings = []
        
        def resolve_file_ref(match):
            link_text = match.group(1)
            file_path = match.group(2)
            
            # Skip URLs and anchors
            if file_path.startswith(('http://', 'https://', 'mailto:', 'ftp://', '#')):
                return match.group(0)
            
            # Resolve relative path
            if not file_path.startswith('/'):
                resolved_path = (source_path.parent / file_path).resolve()
            else:
                resolved_path = Path(file_path)
            
            # Validate file existence if requested
            if validate_links and not resolved_path.exists():
                warnings.append(f"Broken file reference: {file_path}")
                self.broken_links.append({'type': 'file', 'path': file_path})
            
            # Convert to relative path
            try:
                rel_path = os.path.relpath(resolved_path, source_path.parent)
                rel_path = rel_path.replace('\\', '/')
                return f"[{link_text}]({rel_path})"
            except ValueError:
                # If relative path can't be computed, use absolute
                return f"[{link_text}](file://{resolved_path})"
        
        content = re.sub(r'\[([^\]]+)\]\(([^)#]+)\)', resolve_file_ref, content)
        
        return content, warnings
    
    def _resolve_cross_references(self, content: str, source_path: Path) -> Tuple[str, List[str]]:
        """Resolve cross-document references"""
        warnings = []
        
        # Pattern for cross-references like [[Document:Section]]
        def resolve_cross_ref(match):
            ref_full = match.group(1)
            parts = ref_full.split(':', 1)
            
            if len(parts) == 2:
                doc_ref, section_ref = parts
                
                # Look up document in reference map
                if doc_ref in self.reference_map:
                    target_path = self.reference_map[doc_ref]
                    anchor = re.sub(r'[^\w\-]', '-', section_ref.lower()).strip('-')
                    
                    try:
                        rel_path = os.path.relpath(target_path, source_path.parent)
                        rel_path = rel_path.replace('\\', '/')
                        return f"[{ref_full}]({rel_path}#{anchor})"
                    except ValueError:
                        return f"[{ref_full}](file://{target_path}#{anchor})"
                else:
                    warnings.append(f"Unknown document reference: {doc_ref}")
                    self.broken_links.append({'type': 'cross_reference', 'document': doc_ref})
            
            return match.group(0)
        
        content = re.sub(r'\[\[([^\]]+)\]\]', resolve_cross_ref, content)
        
        return content, warnings
    
    def build_reference_map(self, content_files: List[Path]) -> None:
        """Build a map of document titles to file paths"""
        for file_path in content_files:
            if file_path.suffix in ['.md', '.markdown']:
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        first_line = f.readline().strip()
                        
                    # Extract title from first header
                    if first_line.startswith('#'):
                        title = first_line.lstrip('#').strip()
                        # Use title as key
                        self.reference_map[title] = file_path
                        # Also use filename without extension
                        self.reference_map[file_path.stem] = file_path
                except Exception:
                    pass
    
    def get_broken_links_report(self) -> Dict[str, Any]:
        """Get a report of all broken links found"""
        return {
            'total_broken': len(self.broken_links),
            'by_type': self._group_broken_links_by_type(),
            'details': self.broken_links
        }
    
    def _group_broken_links_by_type(self) -> Dict[str, int]:
        """Group broken links by type"""
        type_counts = {}
        for link in self.broken_links:
            link_type = link.get('type', 'unknown')
            type_counts[link_type] = type_counts.get(link_type, 0) + 1
        return type_counts


class ExportTemplate(ABC):
    """Base class for export templates"""
    
    def __init__(self, name: str, base_dir: Path):
        self.name = name
        self.base_dir = base_dir
        self.template_dir = base_dir / "export" / "templates" / name
        self.template_dir.mkdir(parents=True, exist_ok=True)
        
    @abstractmethod
    def get_css_styles(self) -> str:
        """Get CSS styles for this template"""
        pass
        
    @abstractmethod
    def get_html_template(self) -> str:
        """Get HTML template"""
        pass
        
    @abstractmethod
    def process_content(self, content: str, metadata: Dict[str, Any]) -> str:
        """Process content for this template"""
        pass
    
    def get_print_styles(self) -> str:
        """Get print-specific CSS styles"""
        return """
        @media print {
            body {
                font-size: 10pt;
                line-height: 1.4;
            }
            .no-print {
                display: none !important;
            }
            .page-break {
                page-break-before: always;
            }
            a {
                text-decoration: none;
                color: black;
            }
            pre {
                white-space: pre-wrap;
                word-wrap: break-word;
            }
        }
        """


class AcademicTemplate(ExportTemplate):
    """Enhanced academic template with professional styling"""
    
    def get_css_styles(self) -> str:
        base_styles = """
        @page {
            size: A4;
            margin: 2cm 2.5cm 2cm 2.5cm;
            @top-center {
                content: string(document-title);
                font-size: 10pt;
                color: #666;
            }
            @bottom-center {
                content: "Page " counter(page) " of " counter(pages);
                font-size: 10pt;
                color: #666;
            }
        }
        
        :root {
            --primary-color: #1a1a1a;
            --secondary-color: #2c3e50;
            --accent-color: #3498db;
            --text-color: #333;
            --background-color: #fff;
            --code-background: #f8f9fa;
            --border-color: #dee2e6;
        }
        
        body {
            font-family: 'Crimson Text', 'Times New Roman', serif;
            font-size: 11pt;
            line-height: 1.6;
            color: var(--text-color);
            background-color: var(--background-color);
            margin: 0;
            padding: 0;
            text-align: justify;
            hyphens: auto;
        }
        
        /* Typography */
        h1, h2, h3, h4, h5, h6 {
            font-family: 'Montserrat', 'Arial', sans-serif;
            font-weight: 600;
            line-height: 1.2;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
            page-break-after: avoid;
        }
        
        h1 {
            string-set: document-title content();
            color: var(--primary-color);
            font-size: 24pt;
            margin-top: 0;
            padding-bottom: 0.5em;
            border-bottom: 3px solid var(--accent-color);
        }
        
        h2 {
            color: var(--secondary-color);
            font-size: 18pt;
            margin-top: 2em;
            border-bottom: 1px solid var(--border-color);
            padding-bottom: 0.3em;
        }
        
        h3 {
            color: var(--secondary-color);
            font-size: 14pt;
            margin-top: 1.5em;
        }
        
        h4 {
            font-size: 12pt;
            margin-top: 1.2em;
        }
        
        h5, h6 {
            font-size: 11pt;
            margin-top: 1em;
        }
        
        /* Paragraphs and Text */
        p {
            margin: 0 0 1em 0;
            orphans: 3;
            widows: 3;
        }
        
        strong, b {
            font-weight: 600;
        }
        
        em, i {
            font-style: italic;
        }
        
        /* Lists */
        ul, ol {
            margin: 1em 0;
            padding-left: 2em;
        }
        
        li {
            margin: 0.3em 0;
        }
        
        /* Nested lists */
        li > ul, li > ol {
            margin: 0.3em 0;
        }
        
        /* Definition lists */
        dl {
            margin: 1em 0;
        }
        
        dt {
            font-weight: 600;
            margin-top: 0.5em;
        }
        
        dd {
            margin-left: 2em;
            margin-bottom: 0.5em;
        }
        
        /* Links */
        a {
            color: var(--accent-color);
            text-decoration: none;
        }
        
        a:hover {
            text-decoration: underline;
        }
        
        /* Images */
        img {
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1.5em auto;
            page-break-inside: avoid;
            border: 1px solid var(--border-color);
            border-radius: 4px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        
        figure {
            margin: 1.5em 0;
            text-align: center;
            page-break-inside: avoid;
        }
        
        figcaption {
            font-size: 9pt;
            color: #666;
            font-style: italic;
            margin-top: 0.5em;
        }
        
        /* Tables */
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 1.5em 0;
            font-size: 10pt;
            page-break-inside: avoid;
        }
        
        th, td {
            border: 1px solid var(--border-color);
            padding: 8pt;
            text-align: left;
            vertical-align: top;
        }
        
        th {
            background-color: #f8f9fa;
            font-weight: 600;
            color: var(--secondary-color);
        }
        
        tr:nth-child(even) {
            background-color: #f8f9fa;
        }
        
        /* Code blocks */
        code {
            font-family: 'Fira Code', 'Consolas', monospace;
            font-size: 9pt;
            background-color: var(--code-background);
            padding: 2px 4px;
            border-radius: 3px;
            color: #d63384;
        }
        
        pre {
            background-color: var(--code-background);
            border: 1px solid var(--border-color);
            border-radius: 4px;
            padding: 1em;
            overflow-x: auto;
            page-break-inside: avoid;
            margin: 1em 0;
        }
        
        pre code {
            background-color: transparent;
            padding: 0;
            color: inherit;
        }
        
        /* Blockquotes */
        blockquote {
            margin: 1.5em 0;
            padding: 1em 1.5em;
            border-left: 4px solid var(--accent-color);
            background-color: #f8f9fa;
            font-style: italic;
            page-break-inside: avoid;
        }
        
        blockquote cite {
            display: block;
            margin-top: 0.5em;
            font-size: 9pt;
            font-style: normal;
            text-align: right;
        }
        
        /* Horizontal rules */
        hr {
            border: none;
            border-top: 1px solid var(--border-color);
            margin: 2em 0;
        }
        
        /* Special sections */
        .abstract {
            background-color: #f8f9fa;
            border: 1px solid var(--border-color);
            border-radius: 4px;
            padding: 1.5em;
            margin: 2em 0;
            font-size: 10pt;
        }
        
        .abstract h2 {
            font-size: 14pt;
            margin-top: 0;
            border: none;
        }
        
        .keywords {
            margin: 1em 0;
            font-size: 10pt;
        }
        
        .keywords strong {
            color: var(--secondary-color);
        }
        
        /* Table of contents */
        .toc {
            background-color: #f8f9fa;
            border: 1px solid var(--border-color);
            border-radius: 4px;
            padding: 1.5em;
            margin: 2em 0;
            page-break-after: always;
        }
        
        .toc h2 {
            margin-top: 0;
            border: none;
        }
        
        .toc ul {
            list-style: none;
            padding-left: 0;
        }
        
        .toc li {
            margin: 0.5em 0;
            position: relative;
            padding-left: 2em;
        }
        
        .toc li::before {
            content: counter(toc-counter) ". ";
            counter-increment: toc-counter;
            position: absolute;
            left: 0;
            font-weight: 600;
        }
        
        .toc a {
            color: var(--text-color);
            text-decoration: none;
            display: flex;
            justify-content: space-between;
        }
        
        .toc a::after {
            content: leader(".") " " target-counter(attr(href), page);
        }
        
        /* Footnotes */
        .footnotes {
            font-size: 9pt;
            margin-top: 3em;
            padding-top: 1em;
            border-top: 1px solid var(--border-color);
        }
        
        .footnote-ref {
            font-size: 8pt;
            vertical-align: super;
        }
        
        /* Equations */
        .equation {
            text-align: center;
            margin: 1.5em 0;
            page-break-inside: avoid;
        }
        
        .equation-number {
            float: right;
            color: #666;
        }
        
        /* Alerts and callouts */
        .alert {
            padding: 1em;
            margin: 1em 0;
            border-radius: 4px;
            page-break-inside: avoid;
        }
        
        .alert-info {
            background-color: #d1ecf1;
            border: 1px solid #bee5eb;
            color: #0c5460;
        }
        
        .alert-warning {
            background-color: #fff3cd;
            border: 1px solid #ffeaa7;
            color: #856404;
        }
        
        .alert-danger {
            background-color: #f8d7da;
            border: 1px solid #f5c6cb;
            color: #721c24;
        }
        
        .alert-success {
            background-color: #d4edda;
            border: 1px solid #c3e6cb;
            color: #155724;
        }
        
        /* Page elements */
        .page-break {
            page-break-before: always;
        }
        
        .no-break {
            page-break-inside: avoid;
        }
        
        /* Watermark */
        .watermark {
            position: fixed;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%) rotate(-45deg);
            font-size: 72pt;
            color: rgba(0, 0, 0, 0.05);
            z-index: -1;
            font-weight: bold;
        }
        """
        
        return base_styles + self.get_print_styles()
    
    def get_html_template(self) -> str:
        return """
        <!DOCTYPE html>
        <html lang="{lang}">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <meta name="author" content="{author}">
            <meta name="description" content="{description}">
            <meta name="keywords" content="{keywords}">
            <title>{title}</title>
            <link href="https://fonts.googleapis.com/css2?family=Crimson+Text:ital,wght@0,400;0,600;1,400&family=Montserrat:wght@400;600;700&family=Fira+Code:wght@400;500&display=swap" rel="stylesheet">
            <style>{css}</style>
        </head>
        <body>
            {watermark}
            <div class="document">
                {header}
                {content}
                {footer}
            </div>
        </body>
        </html>
        """
    
    def process_content(self, content: str, metadata: Dict[str, Any]) -> str:
        """Process content for academic template"""
        
        # Add metadata section if available
        if metadata:
            meta_section = self._create_metadata_section(metadata)
            content = meta_section + "\n\n" + content
        
        # Process special sections
        content = self._process_abstract(content)
        content = self._process_keywords(content)
        content = self._process_references(content)
        content = self._process_equations(content)
        content = self._process_alerts(content)
        content = self._process_figures(content)
        
        # Add table of contents if requested
        if metadata.get('include_toc', True):
            content = self._add_table_of_contents(content)
        
        return content
    
    def _create_metadata_section(self, metadata: Dict[str, Any]) -> str:
        """Create metadata section"""
        parts = []
        
        if metadata.get('title'):
            parts.append(f"# {metadata['title']}")
        
        meta_items = []
        for key in ['author', 'institution', 'date', 'subject', 'course']:
            if metadata.get(key):
                meta_items.append(f"**{key.title()}:** {metadata[key]}")
        
        if meta_items:
            parts.append("\n".join(meta_items))
        
        return "\n\n".join(parts)
    
    def _process_abstract(self, content: str) -> str:
        """Process abstract sections"""
        content = re.sub(
            r'^## Abstract\s*\n(.*?)(?=\n## |\n# |\Z)',
            r'<div class="abstract">\n<h2>Abstract</h2>\n\1\n</div>',
            content,
            flags=re.MULTILINE | re.DOTALL
        )
        return content
    
    def _process_keywords(self, content: str) -> str:
        """Process keywords sections"""
        content = re.sub(
            r'^\*\*Keywords:\*\*(.+)$',
            r'<div class="keywords"><strong>Keywords:</strong>\1</div>',
            content,
            flags=re.MULTILINE
        )
        return content
    
    def _process_references(self, content: str) -> str:
        """Process references section"""
        # Mark references section for special formatting
        content = re.sub(
            r'^(#+ References)$',
            r'<div class="references">\n\1',
            content,
            flags=re.MULTILINE
        )
        
        # Close references div at end of document or next major section
        if '<div class="references">' in content:
            # Find the end of references section
            parts = content.split('<div class="references">')
            if len(parts) > 1:
                # Find next major section or end of content
                ref_content = parts[1]
                next_section = re.search(r'\n# ', ref_content)
                if next_section:
                    insert_pos = next_section.start()
                    parts[1] = ref_content[:insert_pos] + '\n</div>' + ref_content[insert_pos:]
                else:
                    parts[1] = ref_content + '\n</div>'
                content = '<div class="references">'.join(parts)
        
        # Format reference items
        content = re.sub(
            r'^(\d+\.\s+.+)$',
            r'<div class="reference-item">\1</div>',
            content,
            flags=re.MULTILINE
        )
        
        return content
    
    def _process_equations(self, content: str) -> str:
        """Process equations with numbering"""
        equation_counter = 0
        
        def replace_equation(match):
            nonlocal equation_counter
            equation_counter += 1
            equation_content = match.group(1)
            return f'<div class="equation">\n{equation_content}\n<span class="equation-number">({equation_counter})</span>\n</div>'
        
        # Process display equations
        content = re.sub(
            r'\$\$([^$]+)\$\$',
            replace_equation,
            content,
            flags=re.MULTILINE
        )
        
        return content
    
    def _process_alerts(self, content: str) -> str:
        """Process alert/callout blocks"""
        # Process different types of alerts
        alert_types = {
            'NOTE': 'info',
            'WARNING': 'warning',
            'DANGER': 'danger',
            'SUCCESS': 'success',
            'TIP': 'info',
            'IMPORTANT': 'warning'
        }
        
        for alert_key, alert_class in alert_types.items():
            pattern = rf'^> \*\*{alert_key}:\*\*(.+?)(?=\n(?!>)|\Z)'
            replacement = rf'<div class="alert alert-{alert_class}">\1</div>'
            content = re.sub(pattern, replacement, content, flags=re.MULTILINE | re.DOTALL)
        
        return content
    
    def _process_figures(self, content: str) -> str:
        """Process figures with captions"""
        # Pattern for images with captions
        pattern = r'!\[([^\]]*)\]\(([^)]+)\)\s*\n\s*\*([^*]+)\*'
        
        def replace_figure(match):
            alt_text = match.group(1)
            img_path = match.group(2)
            caption = match.group(3)
            return f'<figure>\n<img src="{img_path}" alt="{alt_text}">\n<figcaption>{caption}</figcaption>\n</figure>'
        
        content = re.sub(pattern, replace_figure, content)
        
        return content
    
    def _add_table_of_contents(self, content: str) -> str:
        """Add table of contents"""
        # Extract headers
        headers = re.findall(r'^(#{1,3})\s+(.+)$', content, re.MULTILINE)
        
        if not headers:
            return content
        
        # Build TOC
        toc_lines = ['<div class="toc">', '<h2>Table of Contents</h2>', '<ul>']
        
        for level, title in headers:
            # Skip the main title and TOC itself
            if title == 'Table of Contents' or (level == '#' and headers.index((level, title)) == 0):
                continue
            
            depth = len(level) - 1
            indent = '  ' * depth
            anchor = re.sub(r'[^\w\-]', '-', title.lower()).strip('-')
            toc_lines.append(f'{indent}<li><a href="#{anchor}">{title}</a></li>')
        
        toc_lines.extend(['</ul>', '</div>', ''])
        
        # Insert TOC after title/metadata
        lines = content.split('\n')
        insert_pos = 0
        
        # Find position after title and metadata
        for i, line in enumerate(lines):
            if line.startswith('#') and i > 0:
                insert_pos = i
                break
            elif not line.strip() and i > 2:
                insert_pos = i
                break
        
        lines[insert_pos:insert_pos] = toc_lines
        return '\n'.join(lines)


class PDFExporter:
    """Enhanced PDF exporter with multiple backend support"""
    
    def __init__(self, template: ExportTemplate):
        self.template = template
        self.backend = self._determine_backend()
        
    def _determine_backend(self) -> str:
        """Determine which PDF backend to use"""
        if WEASYPRINT_AVAILABLE:
            return 'weasyprint'
        elif REPORTLAB_AVAILABLE:
            return 'reportlab'
        else:
            return 'none'
    
    def export(self, content: str, output_path: Path, 
              config: ExportConfig, metadata: Dict[str, Any] = None) -> ExportResult:
        """Export content to PDF"""
        
        start_time = time.time()
        errors = []
        warnings = []
        
        try:
            # Process content with template
            processed_content = self.template.process_content(content, metadata or {})
            
            # Choose export method based on backend
            if self.backend == 'weasyprint':
                result = self._export_with_weasyprint(
                    processed_content, output_path, config, metadata
                )
            elif self.backend == 'reportlab':
                result = self._export_with_reportlab(
                    processed_content, output_path, config, metadata
                )
            else:
                # Fallback to HTML export with print styles
                result = self._export_html_as_pdf(
                    processed_content, output_path, config, metadata
                )
            
            if result['success']:
                file_size = output_path.stat().st_size / (1024 * 1024)
                validation_score = self._validate_pdf(output_path)
                
                # Generate quality report
                quality_report = self._generate_quality_report(output_path)
                
                return ExportResult(
                    success=True,
                    output_files=[str(output_path)],
                    format_type='pdf',
                    file_size_mb=file_size,
                    processing_time=time.time() - start_time,
                    validation_score=validation_score,
                    errors=errors,
                    warnings=warnings + result.get('warnings', []),
                    metadata={
                        'pages': result.get('pages', 0),
                        'backend': self.backend,
                        'compression_ratio': result.get('compression_ratio', 0)
                    },
                    quality_report=quality_report
                )
            else:
                errors.append(result.get('error', 'PDF generation failed'))
                
        except Exception as e:
            errors.append(f"PDF export failed: {str(e)}")
        
        return ExportResult(
            success=False,
            output_files=[],
            format_type='pdf',
            file_size_mb=0,
            processing_time=time.time() - start_time,
            validation_score=0,
            errors=errors,
            warnings=warnings,
            metadata={}
        )
    
    def _export_with_weasyprint(self, content: str, output_path: Path,
                               config: ExportConfig, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Export using WeasyPrint"""
        try:
            # Convert markdown to HTML
            html_content = self._markdown_to_html(content, enhanced=True)
            
            # Prepare metadata
            template_metadata = {
                'title': metadata.get('title', 'Document'),
                'author': metadata.get('author', ''),
                'description': metadata.get('description', ''),
                'keywords': metadata.get('keywords', ''),
                'lang': metadata.get('language', 'en')
            }
            
            # Add watermark if requested
            watermark_html = ''
            if config.watermark_text:
                watermark_html = f'<div class="watermark">{config.watermark_text}</div>'
            
            # Add header/footer if requested
            header_html = ''
            footer_html = ''
            if config.header_text:
                header_html = f'<div class="header">{config.header_text}</div>'
            if config.footer_text:
                footer_html = f'<div class="footer">{config.footer_text}</div>'
            
            # Create full HTML document
            full_html = self.template.get_html_template().format(
                **template_metadata,
                css=self.template.get_css_styles(),
                content=html_content,
                watermark=watermark_html,
                header=header_html,
                footer=footer_html
            )
            
            # Configure fonts
            font_config = FontConfiguration()
            
            # Add custom CSS if provided
            css_list = []
            if config.custom_css:
                css_list.append(CSS(string=config.custom_css, font_config=font_config))
            
            # Generate PDF
            doc = HTML(string=full_html, base_url=str(output_path.parent))
            
            # Apply optimization settings
            pdf_options = {}
            if config.enable_compression:
                pdf_options['optimize_size'] = ('fonts', 'images')
            
            doc.write_pdf(
                output_path,
                stylesheets=css_list,
                font_config=font_config,
                **pdf_options
            )
            
            # Get page count
            page_count = self._count_pdf_pages(output_path)
            
            return {
                'success': True,
                'pages': page_count,
                'warnings': []
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f"WeasyPrint error: {str(e)}",
                'warnings': []
            }
    
    def _export_with_reportlab(self, content: str, output_path: Path,
                              config: ExportConfig, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Export using ReportLab"""
        try:
            from reportlab.lib.pagesizes import A4, letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.platypus import Table, TableStyle, Image as RLImage
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
            from reportlab.lib.units import inch
            
            # Set page size
            page_size = A4 if config.page_size == 'A4' else letter
            
            # Create document
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=page_size,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72,
                title=metadata.get('title', 'Document'),
                author=metadata.get('author', ''),
                subject=metadata.get('subject', '')
            )
            
            # Build story
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            styles.add(ParagraphStyle(
                name='Justify',
                parent=styles['Normal'],
                alignment=TA_JUSTIFY,
                fontSize=11,
                leading=14
            ))
            
            # Parse markdown and build PDF elements
            elements = self._markdown_to_reportlab_elements(content, styles)
            story.extend(elements)
            
            # Build PDF
            doc.build(story)
            
            # Get page count
            page_count = self._count_pdf_pages(output_path)
            
            return {
                'success': True,
                'pages': page_count,
                'warnings': []
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f"ReportLab error: {str(e)}",
                'warnings': []
            }
    
    def _export_html_as_pdf(self, content: str, output_path: Path,
                           config: ExportConfig, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Export HTML with print styles as fallback"""
        try:
            # Convert markdown to HTML
            html_content = self._markdown_to_html(content, enhanced=True)
            
            # Create HTML file with print styles
            html_path = output_path.with_suffix('.html')
            
            template_metadata = {
                'title': metadata.get('title', 'Document'),
                'author': metadata.get('author', ''),
                'description': metadata.get('description', ''),
                'keywords': metadata.get('keywords', ''),
                'lang': metadata.get('language', 'en'),
                'css': self.template.get_css_styles(),
                'content': html_content,
                'watermark': '',
                'header': '',
                'footer': ''
            }
            
            full_html = self.template.get_html_template().format(**template_metadata)
            
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
            
            # Try to use system tools to convert HTML to PDF
            success = self._convert_html_to_pdf_system(html_path, output_path)
            
            if success:
                # Clean up HTML file
                html_path.unlink()
                return {
                    'success': True,
                    'pages': 0,
                    'warnings': ['PDF generated using HTML export with print styles']
                }
            else:
                return {
                    'success': False,
                    'error': 'No PDF backend available. HTML file created instead.',
                    'warnings': [f'HTML file saved to: {html_path}']
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': f"HTML to PDF conversion error: {str(e)}",
                'warnings': []
            }
    
    def _convert_html_to_pdf_system(self, html_path: Path, pdf_path: Path) -> bool:
        """Try to convert HTML to PDF using system tools"""
        # Try wkhtmltopdf
        try:
            result = subprocess.run(
                ['wkhtmltopdf', '--quiet', str(html_path), str(pdf_path)],
                capture_output=True,
                timeout=60
            )
            if result.returncode == 0:
                return True
        except (subprocess.SubprocessError, FileNotFoundError):
            pass
        
        # Try Chrome/Chromium headless
        for browser in ['google-chrome', 'chromium', 'chromium-browser']:
            try:
                result = subprocess.run(
                    [browser, '--headless', '--print-to-pdf=' + str(pdf_path), str(html_path)],
                    capture_output=True,
                    timeout=60
                )
                if result.returncode == 0:
                    return True
            except (subprocess.SubprocessError, FileNotFoundError):
                continue
        
        return False
    
    def _markdown_to_html(self, content: str, enhanced: bool = False) -> str:
        """Convert markdown to HTML with optional enhancements"""
        if MARKDOWN_AVAILABLE and enhanced:
            # Use full markdown library with extensions
            md = markdown.Markdown(extensions=[
                'extra',
                'tables',
                'fenced_code',
                'codehilite',
                'toc',
                'attr_list',
                'def_list',
                'footnotes',
                'admonition',
                'sane_lists'
            ])
            return md.convert(content)
        else:
            # Fallback to simple conversion
            return self._simple_markdown_to_html(content)
    
    def _simple_markdown_to_html(self, content: str) -> str:
        """Simple markdown to HTML conversion"""
        lines = content.split('\n')
        html_lines = []
        in_code_block = False
        in_list = False
        list_stack = []
        
        for i, line in enumerate(lines):
            # Code blocks
            if line.startswith('```'):
                if in_code_block:
                    html_lines.append('</code></pre>')
                    in_code_block = False
                else:
                    lang = line[3:].strip()
                    html_lines.append(f'<pre><code class="language-{lang}">')
                    in_code_block = True
                continue
            
            if in_code_block:
                html_lines.append(self._escape_html(line))
                continue
            
            # Headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                if level <= 6:
                    text = line[level:].strip()
                    anchor = re.sub(r'[^\w\-]', '-', text.lower()).strip('-')
                    html_lines.append(f'<h{level} id="{anchor}">{self._process_inline(text)}</h{level}>')
                    continue
            
            # Lists
            list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', line)
            if list_match:
                indent_level = len(list_match.group(1)) // 2
                list_type = 'ol' if list_match.group(2)[0].isdigit() else 'ul'
                item_text = list_match.group(3)
                
                # Manage list stack
                while len(list_stack) > indent_level + 1:
                    closed_type = list_stack.pop()
                    html_lines.append(f'</li></{closed_type}>')
                
                if len(list_stack) == indent_level:
                    html_lines.append(f'<{list_type}>')
                    list_stack.append(list_type)
                elif len(list_stack) == indent_level + 1:
                    if list_stack[-1] != list_type:
                        html_lines.append(f'</{list_stack[-1]}><{list_type}>')
                        list_stack[-1] = list_type
                    else:
                        html_lines.append('</li>')
                
                html_lines.append(f'<li>{self._process_inline(item_text)}')
                continue
            
            # Close lists if needed
            if list_stack and not line.strip():
                while list_stack:
                    closed_type = list_stack.pop()
                    html_lines.append(f'</li></{closed_type}>')
            
            # Horizontal rule
            if re.match(r'^[-_*]{3,}\s*$', line):
                html_lines.append('<hr>')
                continue
            
            # Blockquote
            if line.startswith('>'):
                quote_text = line[1:].strip()
                html_lines.append(f'<blockquote>{self._process_inline(quote_text)}</blockquote>')
                continue
            
            # Paragraph
            if line.strip():
                html_lines.append(f'<p>{self._process_inline(line)}</p>')
            else:
                html_lines.append('')
        
        # Close any remaining lists
        while list_stack:
            closed_type = list_stack.pop()
            html_lines.append(f'</li></{closed_type}>')
        
        return '\n'.join(html_lines)
    
    def _process_inline(self, text: str) -> str:
        """Process inline markdown formatting"""
        # Escape HTML
        text = self._escape_html(text)
        
        # Strong
        text = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'__([^_]+)__', r'<strong>\1</strong>', text)
        
        # Emphasis
        text = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', text)
        text = re.sub(r'_([^_]+)_', r'<em>\1</em>', text)
        
        # Code
        text = re.sub(r'`([^`]+)`', r'<code>\1</code>', text)
        
        # Links
        text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', text)
        
        # Images
        text = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', r'<img src="\2" alt="\1">', text)
        
        return text
    
    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters"""
        return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&#39;'))
    
    def _markdown_to_reportlab_elements(self, content: str, styles) -> List:
        """Convert markdown to ReportLab elements"""
        elements = []
        lines = content.split('\n')
        
        for line in lines:
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line[level:].strip()
                style = f'Heading{min(level, 6)}'
                elements.append(Paragraph(text, styles[style]))
                elements.append(Spacer(1, 0.2*inch))
            elif line.strip():
                elements.append(Paragraph(line, styles['Justify']))
                elements.append(Spacer(1, 0.1*inch))
        
        return elements
    
    def _validate_pdf(self, pdf_path: Path) -> float:
        """Validate PDF quality"""
        try:
            if not pdf_path.exists():
                return 0.0
            
            file_size = pdf_path.stat().st_size
            if file_size < 1024:  # Less than 1KB
                return 0.1
            
            # Try to verify it's a valid PDF
            with open(pdf_path, 'rb') as f:
                header = f.read(4)
                if header != b'%PDF':
                    return 0.2
            
            # Try to count pages
            page_count = self._count_pdf_pages(pdf_path)
            if page_count == 0:
                return 0.3
            
            # Basic score based on file size and pages
            score = min(0.8 + (page_count * 0.02), 1.0)
            return score
            
        except Exception:
            return 0.3
    
    def _count_pdf_pages(self, pdf_path: Path) -> int:
        """Count pages in PDF"""
        try:
            # Try using PyPDF2
            import PyPDF2
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return len(reader.pages)
        except ImportError:
            pass
        
        try:
            # Try using pdfplumber
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                return len(pdf.pages)
        except ImportError:
            pass
        
        # Fallback: estimate based on file size
        file_size = pdf_path.stat().st_size
        estimated_pages = max(1, file_size // (50 * 1024))  # ~50KB per page
        return estimated_pages
    
    def _generate_quality_report(self, pdf_path: Path) -> Dict[str, Any]:
        """Generate quality report for PDF"""
        report = {
            'file_valid': pdf_path.exists(),
            'file_size_mb': pdf_path.stat().st_size / (1024 * 1024) if pdf_path.exists() else 0,
            'page_count': self._count_pdf_pages(pdf_path),
            'has_text': True,  # Would need PDF text extraction to verify
            'has_images': True,  # Would need PDF parsing to verify
            'compression_effective': True
        }
        
        # Add quality score
        quality_score = sum([
            0.3 if report['file_valid'] else 0,
            0.2 if report['page_count'] > 0 else 0,
            0.2 if report['file_size_mb'] < 10 else 0.1,
            0.3 if report['has_text'] else 0
        ])
        
        report['quality_score'] = quality_score
        return report


class HTMLExporter:
    """Enhanced HTML exporter with responsive design"""
    
    def __init__(self, template: ExportTemplate):
        self.template = template
        
    def export(self, content: str, output_path: Path,
              config: ExportConfig, metadata: Dict[str, Any] = None) -> ExportResult:
        """Export content to HTML"""
        
        start_time = time.time()
        errors = []
        warnings = []
        
        try:
            # Process content with template
            processed_content = self.template.process_content(content, metadata or {})
            
            # Convert markdown to HTML
            html_content = self._markdown_to_html(processed_content)
            
            # Add responsive design elements
            html_content = self._add_responsive_design(html_content)
            
            # Prepare metadata
            template_metadata = {
                'title': metadata.get('title', 'Document'),
                'author': metadata.get('author', ''),
                'description': metadata.get('description', ''),
                'keywords': metadata.get('keywords', ''),
                'lang': metadata.get('language', 'en')
            }
            
            # Add navigation if requested
            if config.include_toc:
                html_content = self._add_navigation(html_content)
            
            # Create full HTML document
            full_html = self.template.get_html_template().format(
                **template_metadata,
                css=self._get_enhanced_css(),
                content=html_content,
                watermark='',
                header='',
                footer=''
            )
            
            # Write HTML file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
            
            # Validate output
            if output_path.exists():
                file_size = output_path.stat().st_size / (1024 * 1024)
                validation_score = self._validate_html(output_path)
                
                return ExportResult(
                    success=True,
                    output_files=[str(output_path)],
                    format_type='html',
                    file_size_mb=file_size,
                    processing_time=time.time() - start_time,
                    validation_score=validation_score,
                    errors=errors,
                    warnings=warnings,
                    metadata={
                        'word_count': len(content.split()),
                        'has_responsive': True,
                        'has_navigation': config.include_toc
                    }
                )
            else:
                errors.append("HTML file was not created")
                
        except Exception as e:
            errors.append(f"HTML export failed: {str(e)}")
        
        return ExportResult(
            success=False,
            output_files=[],
            format_type='html',
            file_size_mb=0,
            processing_time=time.time() - start_time,
            validation_score=0,
            errors=errors,
            warnings=warnings,
            metadata={}
        )
    
    def _markdown_to_html(self, content: str) -> str:
        """Convert markdown to HTML with enhanced features"""
        if MARKDOWN_AVAILABLE:
            md = markdown.Markdown(extensions=[
                'extra',
                'tables',
                'fenced_code',
                'codehilite',
                'toc',
                'attr_list',
                'def_list',
                'footnotes',
                'admonition',
                'sane_lists',
                'smarty'
            ])
            return md.convert(content)
        else:
            # Use PDF exporter's simple converter
            exporter = PDFExporter(self.template)
            return exporter._simple_markdown_to_html(content)
    
    def _get_enhanced_css(self) -> str:
        """Get enhanced CSS with responsive design"""
        base_css = self.template.get_css_styles()
        
        responsive_css = """
        /* Responsive Design */
        @media screen and (max-width: 768px) {
            body {
                font-size: 14px;
                padding: 10px;
            }
            
            h1 { font-size: 20pt; }
            h2 { font-size: 16pt; }
            h3 { font-size: 14pt; }
            
            table {
                font-size: 12px;
                overflow-x: auto;
                display: block;
            }
            
            img {
                max-width: 100%;
            }
            
            .toc {
                position: static;
                width: 100%;
            }
        }
        
        /* Navigation */
        .nav-container {
            position: fixed;
            top: 0;
            left: 0;
            width: 250px;
            height: 100vh;
            background-color: #f8f9fa;
            border-right: 1px solid #dee2e6;
            overflow-y: auto;
            transition: transform 0.3s;
            z-index: 1000;
        }
        
        .nav-container.hidden {
            transform: translateX(-100%);
        }
        
        .nav-toggle {
            position: fixed;
            top: 10px;
            left: 10px;
            z-index: 1001;
            background-color: #007bff;
            color: white;
            border: none;
            padding: 10px;
            border-radius: 4px;
            cursor: pointer;
        }
        
        .content-container {
            margin-left: 250px;
            transition: margin-left 0.3s;
        }
        
        .content-container.full-width {
            margin-left: 0;
        }
        
        /* Search */
        .search-container {
            padding: 10px;
            border-bottom: 1px solid #dee2e6;
        }
        
        .search-input {
            width: 100%;
            padding: 8px;
            border: 1px solid #dee2e6;
            border-radius: 4px;
        }
        
        .search-results {
            max-height: 200px;
            overflow-y: auto;
            border: 1px solid #dee2e6;
            border-radius: 4px;
            margin-top: 10px;
            display: none;
        }
        
        .search-result-item {
            padding: 8px;
            cursor: pointer;
        }
        
        .search-result-item:hover {
            background-color: #e9ecef;
        }
        
        /* Dark mode */
        @media (prefers-color-scheme: dark) {
            :root {
                --primary-color: #f8f9fa;
                --secondary-color: #adb5bd;
                --accent-color: #4dabf7;
                --text-color: #f8f9fa;
                --background-color: #212529;
                --code-background: #343a40;
                --border-color: #495057;
            }
        }
        """
        
        return base_css + responsive_css
    
    def _add_responsive_design(self, content: str) -> str:
        """Add responsive design elements"""
        # Wrap content in container
        content = f'<div class="content-container">\n{content}\n</div>'
        
        # Add viewport meta tag (handled in template)
        return content
    
    def _add_navigation(self, content: str) -> str:
        """Add navigation sidebar"""
        # Extract headers for navigation
        headers = re.findall(r'<h([1-6])\s+id="([^"]+)">([^<]+)</h[1-6]>', content)
        
        if not headers:
            return content
        
        # Build navigation HTML
        nav_html = ['<div class="nav-container">',
                   '<div class="search-container">',
                   '<input type="text" class="search-input" placeholder="Search...">',
                   '<div class="search-results"></div>',
                   '</div>',
                   '<nav class="navigation">']
        
        current_level = 0
        for level, anchor, title in headers:
            level = int(level)
            
            # Close/open nested lists
            if level > current_level:
                nav_html.append('<ul>' * (level - current_level))
            elif level < current_level:
                nav_html.append('</ul>' * (current_level - level))
            
            nav_html.append(f'<li><a href="#{anchor}">{title}</a></li>')
            current_level = level
        
        # Close remaining lists
        nav_html.append('</ul>' * current_level)
        nav_html.extend(['</nav>', '</div>'])
        
        # Add toggle button
        nav_html.append('<button class="nav-toggle">☰</button>')
        
        # Add JavaScript for interactivity
        nav_html.append(self._get_navigation_script())
        
        return '\n'.join(nav_html) + '\n' + content
    
    def _get_navigation_script(self) -> str:
        """Get JavaScript for navigation functionality"""
        return """
        <script>
        // Navigation toggle
        document.querySelector('.nav-toggle').addEventListener('click', function() {
            document.querySelector('.nav-container').classList.toggle('hidden');
            document.querySelector('.content-container').classList.toggle('full-width');
        });
        
        // Search functionality
        const searchInput = document.querySelector('.search-input');
        const searchResults = document.querySelector('.search-results');
        
        searchInput.addEventListener('input', function() {
            const query = this.value.toLowerCase();
            if (query.length < 2) {
                searchResults.style.display = 'none';
                return;
            }
            
            const results = [];
            const elements = document.querySelectorAll('h1, h2, h3, h4, h5, h6, p');
            
            elements.forEach(function(el) {
                if (el.textContent.toLowerCase().includes(query)) {
                    results.push({
                        text: el.textContent.substring(0, 100) + '...',
                        element: el
                    });
                }
            });
            
            if (results.length > 0) {
                searchResults.innerHTML = results.slice(0, 10).map(function(r) {
                    return '<div class="search-result-item">' + r.text + '</div>';
                }).join('');
                searchResults.style.display = 'block';
                
                // Add click handlers
                searchResults.querySelectorAll('.search-result-item').forEach(function(item, i) {
                    item.addEventListener('click', function() {
                        results[i].element.scrollIntoView({behavior: 'smooth'});
                        searchResults.style.display = 'none';
                        searchInput.value = '';
                    });
                });
            } else {
                searchResults.style.display = 'none';
            }
        });
        
        // Smooth scrolling for navigation links
        document.querySelectorAll('.navigation a').forEach(function(link) {
            link.addEventListener('click', function(e) {
                e.preventDefault();
                const target = document.querySelector(this.getAttribute('href'));
                if (target) {
                    target.scrollIntoView({behavior: 'smooth'});
                }
            });
        });
        </script>
        """
    
    def _validate_html(self, html_path: Path) -> float:
        """Validate HTML quality"""
        try:
            if not html_path.exists():
                return 0.0
            
            with open(html_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            score = 0.5  # Base score
            
            # Check for basic HTML structure
            if '<html' in content and '</html>' in content:
                score += 0.2
            if '<head>' in content and '</head>' in content:
                score += 0.1
            if '<body>' in content and '</body>' in content:
                score += 0.1
            if '<title>' in content:
                score += 0.1
            
            return min(score, 1.0)
            
        except Exception:
            return 0.3


class DOCXExporter:
    """Enhanced DOCX exporter with advanced formatting"""
    
    def __init__(self, template: ExportTemplate):
        self.template = template
        
    def export(self, content: str, output_path: Path,
              config: ExportConfig, metadata: Dict[str, Any] = None) -> ExportResult:
        """Export content to DOCX"""
        
        start_time = time.time()
        errors = []
        warnings = []
        
        if not PYTHON_DOCX_AVAILABLE:
            errors.append("python-docx not available, DOCX export failed")
            return ExportResult(
                success=False,
                output_files=[],
                format_type='docx',
                file_size_mb=0,
                processing_time=time.time() - start_time,
                validation_score=0,
                errors=errors,
                warnings=warnings,
                metadata={}
            )
        
        try:
            # Create document
            doc = Document()
            
            # Set document properties
            if metadata:
                self._set_document_properties(doc, metadata)
            
            # Process content
            processed_content = self.template.process_content(content, metadata or {})
            
            # Apply styles
            self._setup_styles(doc)
            
            # Parse markdown and add to document
            self._markdown_to_docx(processed_content, doc, config)
            
            # Add table of contents if requested
            if config.include_toc:
                self._add_table_of_contents(doc)
            
            # Save document
            doc.save(output_path)
            
            if output_path.exists():
                file_size = output_path.stat().st_size / (1024 * 1024)
                validation_score = self._validate_docx(output_path)
                
                return ExportResult(
                    success=True,
                    output_files=[str(output_path)],
                    format_type='docx',
                    file_size_mb=file_size,
                    processing_time=time.time() - start_time,
                    validation_score=validation_score,
                    errors=errors,
                    warnings=warnings,
                    metadata={
                        'paragraphs': len(doc.paragraphs),
                        'sections': len(doc.sections),
                        'tables': len(doc.tables)
                    }
                )
            else:
                errors.append("DOCX file was not created")
                
        except Exception as e:
            errors.append(f"DOCX export failed: {str(e)}")
        
        return ExportResult(
            success=False,
            output_files=[],
            format_type='docx',
            file_size_mb=0,
            processing_time=time.time() - start_time,
            validation_score=0,
            errors=errors,
            warnings=warnings,
            metadata={}
        )
    
    def _set_document_properties(self, doc: Document, metadata: Dict[str, Any]) -> None:
        """Set document properties"""
        props = doc.core_properties
        
        if metadata.get('title'):
            props.title = metadata['title']
        if metadata.get('author'):
            props.author = metadata['author']
        if metadata.get('subject'):
            props.subject = metadata['subject']
        if metadata.get('keywords'):
            props.keywords = metadata['keywords']
        if metadata.get('description'):
            props.comments = metadata['description']
        
        props.created = datetime.now()
        props.modified = datetime.now()
    
    def _setup_styles(self, doc: Document) -> None:
        """Setup document styles"""
        # Get or create styles
        styles = doc.styles
        
        # Modify normal style
        normal_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.line_spacing = 1.5
        normal_style.paragraph_format.space_after = Pt(6)
        
        # Create custom styles
        if 'Academic Body' not in styles:
            body_style = styles.add_style('Academic Body', WD_STYLE_TYPE.PARAGRAPH)
            body_style.base_style = styles['Normal']
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_style.paragraph_format.first_line_indent = Inches(0.5)
        
        # Heading styles
        for i in range(1, 4):
            heading_style = styles[f'Heading {i}']
            heading_style.font.name = 'Arial'
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
    
    def _markdown_to_docx(self, content: str, doc: Document, config: ExportConfig) -> None:
        """Convert markdown content to DOCX document"""
        lines = content.split('\n')
        current_list = None
        code_block = []
        in_code_block = False
        
        for line in lines:
            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    self._add_code_block(doc, '\n'.join(code_block))
                    code_block = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
                continue
            
            if in_code_block:
                code_block.append(line)
                continue
            
            # Headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                if level <= 9:
                    text = line[level:].strip()
                    doc.add_heading(text, level=min(level, 9))
                continue
            
            # Lists
            list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', line)
            if list_match:
                indent_level = len(list_match.group(1)) // 2
                is_numbered = list_match.group(2)[0].isdigit()
                item_text = list_match.group(3)
                
                # Add list item
                p = doc.add_paragraph(item_text)
                if is_numbered:
                    p.style = 'List Number'
                else:
                    p.style = 'List Bullet'
                
                # Set indentation
                p.paragraph_format.left_indent = Inches(0.5 * indent_level)
                continue
            
            # Images
            img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if img_match:
                alt_text = img_match.group(1)
                img_path = img_match.group(2)
                self._add_image(doc, img_path, alt_text, config)
                continue
            
            # Tables (simple support)
            if '|' in line and line.strip().startswith('|'):
                # Collect table lines
                table_lines = [line]
                # This is simplified - would need more complex parsing
                self._add_simple_table(doc, table_lines)
                continue
            
            # Regular paragraph
            if line.strip():
                p = doc.add_paragraph(line)
                p.style = 'Academic Body'
            else:
                # Empty line - add paragraph break
                doc.add_paragraph()
    
    def _add_code_block(self, doc: Document, code: str) -> None:
        """Add code block to document"""
        p = doc.add_paragraph()
        p.style = 'Normal'
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        # Add code with monospace font
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    def _add_image(self, doc: Document, img_path: str, alt_text: str, config: ExportConfig) -> None:
        """Add image to document"""
        try:
            # Skip URLs
            if img_path.startswith(('http://', 'https://')):
                doc.add_paragraph(f"[Image: {alt_text}] ({img_path})")
                return
            
            # Try to add the image
            if Path(img_path).exists():
                # Determine width based on config
                width_map = {
                    'small': 3,
                    'medium': 4.5,
                    'large': 6,
                    'original': None
                }
                width = width_map.get(config.image_sizing, 4.5)
                
                if width:
                    doc.add_picture(img_path, width=Inches(width))
                else:
                    doc.add_picture(img_path)
                
                # Add caption if alt text provided
                if alt_text:
                    caption = doc.add_paragraph(alt_text)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.style = 'Caption'
            else:
                doc.add_paragraph(f"[Missing image: {alt_text}]")
                
        except Exception as e:
            doc.add_paragraph(f"[Error loading image: {alt_text}]")
    
    def _add_simple_table(self, doc: Document, table_lines: List[str]) -> None:
        """Add simple table to document"""
        # This is a simplified implementation
        # Would need more complex parsing for full table support
        rows = []
        for line in table_lines:
            if line.strip() and '|' in line:
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if cells:
                    rows.append(cells)
        
        if rows:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Table Grid'
            
            for i, row_data in enumerate(rows):
                for j, cell_data in enumerate(row_data[:len(table.columns)]):
                    table.cell(i, j).text = cell_data
    
    def _add_table_of_contents(self, doc: Document) -> None:
        """Add table of contents"""
        # Note: python-docx doesn't support automatic TOC generation
        # This adds a placeholder
        doc.add_page_break()
        doc.add_heading('Table of Contents', level=1)
        doc.add_paragraph(
            'Note: Please update the table of contents using your word processor\'s '
            'table of contents feature.'
        )
        doc.add_page_break()
    
    def _validate_docx(self, docx_path: Path) -> float:
        """Validate DOCX quality"""
        try:
            if not docx_path.exists():
                return 0.0
            
            # Basic validation
            file_size = docx_path.stat().st_size
            if file_size < 1024:  # Less than 1KB
                return 0.1
            
            # Try to open and check basic structure
            try:
                doc = Document(docx_path)
                
                if len(doc.paragraphs) > 0:
                    return 0.8
                else:
                    return 0.3
            except Exception:
                return 0.2
            
        except Exception:
            return 0.2


class EPUBExporter:
    """EPUB exporter for e-book format"""
    
    def __init__(self, template: ExportTemplate):
        self.template = template
        
    def export(self, content: str, output_path: Path,
              config: ExportConfig, metadata: Dict[str, Any] = None) -> ExportResult:
        """Export content to EPUB"""
        
        start_time = time.time()
        errors = []
        warnings = []
        
        if not EBOOKLIB_AVAILABLE:
            errors.append("ebooklib not available, EPUB export failed")
            return ExportResult(
                success=False,
                output_files=[],
                format_type='epub',
                file_size_mb=0,
                processing_time=time.time() - start_time,
                validation_score=0,
                errors=errors,
                warnings=warnings,
                metadata={}
            )
        
        try:
            # Create EPUB book
            book = epub.EpubBook()
            
            # Set metadata
            book.set_identifier(str(uuid.uuid4()))
            book.set_title(metadata.get('title', 'Document'))
            book.set_language(metadata.get('language', 'en'))
            
            if metadata.get('author'):
                book.add_author(metadata['author'])
            
            # Process content
            processed_content = self.template.process_content(content, metadata or {})
            
            # Convert to HTML
            html_content = self._markdown_to_html(processed_content)
            
            # Create chapter
            chapter = epub.EpubHtml(
                title=metadata.get('title', 'Chapter 1'),
                file_name='chapter1.xhtml',
                lang=metadata.get('language', 'en')
            )
            chapter.content = html_content
            
            # Add chapter to book
            book.add_item(chapter)
            
            # Add navigation
            book.toc = (epub.Link('chapter1.xhtml', metadata.get('title', 'Chapter 1'), 'chapter1'),)
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            # Create spine
            book.spine = ['nav', chapter]
            
            # Write EPUB file
            epub.write_epub(str(output_path), book, {})
            
            if output_path.exists():
                file_size = output_path.stat().st_size / (1024 * 1024)
                validation_score = self._validate_epub(output_path)
                
                return ExportResult(
                    success=True,
                    output_files=[str(output_path)],
                    format_type='epub',
                    file_size_mb=file_size,
                    processing_time=time.time() - start_time,
                    validation_score=validation_score,
                    errors=errors,
                    warnings=warnings,
                    metadata={'chapters': 1}
                )
            else:
                errors.append("EPUB file was not created")
                
        except Exception as e:
            errors.append(f"EPUB export failed: {str(e)}")
        
        return ExportResult(
            success=False,
            output_files=[],
            format_type='epub',
            file_size_mb=0,
            processing_time=time.time() - start_time,
            validation_score=0,
            errors=errors,
            warnings=warnings,
            metadata={}
        )
    
    def _markdown_to_html(self, content: str) -> str:
        """Convert markdown to HTML for EPUB"""
        # Reuse HTML exporter's converter
        html_exporter = HTMLExporter(self.template)
        return html_exporter._markdown_to_html(content)
    
    def _validate_epub(self, epub_path: Path) -> float:
        """Validate EPUB file"""
        try:
            if not epub_path.exists():
                return 0.0
            
            # Check if it's a valid ZIP file (EPUB is ZIP-based)
            try:
                with zipfile.ZipFile(epub_path, 'r') as z:
                    # Check for required files
                    files = z.namelist()
                    has_mimetype = 'mimetype' in files
                    has_container = 'META-INF/container.xml' in files
                    
                    if has_mimetype and has_container:
                        return 0.8
                    else:
                        return 0.4
            except zipfile.BadZipFile:
                return 0.2
                
        except Exception:
            return 0.2


import uuid


class ExportQualityValidator:
    """Comprehensive export quality validation system"""
    
    def __init__(self):
        self.validation_rules = {
            'content_completeness': {
                'weight': 0.3,
                'description': 'Checks if all content is properly exported'
            },
            'formatting_preservation': {
                'weight': 0.25,
                'description': 'Validates that formatting is preserved'
            },
            'image_handling': {
                'weight': 0.15,
                'description': 'Ensures images are properly handled'
            },
            'reference_integrity': {
                'weight': 0.15,
                'description': 'Validates reference links and citations'
            },
            'file_validity': {
                'weight': 0.15,
                'description': 'Checks output file validity and structure'
            }
        }
    
    def validate_export(self, export_result: ExportResult, 
                       original_content: str,
                       export_path: Path) -> Dict[str, Any]:
        """Perform comprehensive validation of export result"""
        
        validation_report = {
            'overall_score': 0.0,
            'rule_scores': {},
            'issues': [],
            'recommendations': [],
            'validation_time': 0,
            'file_analysis': {}
        }
        
        start_time = time.time()
        
        try:
            # Content completeness validation
            content_score = self._validate_content_completeness(
                original_content, export_path, export_result.format_type
            )
            validation_report['rule_scores']['content_completeness'] = content_score
            
            # Formatting preservation validation
            format_score = self._validate_formatting_preservation(
                original_content, export_path, export_result.format_type
            )
            validation_report['rule_scores']['formatting_preservation'] = format_score
            
            # Image handling validation
            image_score = self._validate_image_handling(
                original_content, export_path, export_result.format_type
            )
            validation_report['rule_scores']['image_handling'] = image_score
            
            # Reference integrity validation
            ref_score = self._validate_reference_integrity(
                original_content, export_path, export_result.format_type
            )
            validation_report['rule_scores']['reference_integrity'] = ref_score
            
            # File validity validation
            file_score = self._validate_file_validity(
                export_path, export_result.format_type
            )
            validation_report['rule_scores']['file_validity'] = file_score
            
            # Calculate overall score
            overall_score = 0
            for rule, score in validation_report['rule_scores'].items():
                weight = self.validation_rules[rule]['weight']
                overall_score += score * weight
            
            validation_report['overall_score'] = overall_score
            validation_report['validation_time'] = time.time() - start_time
            
            # Generate recommendations
            validation_report['recommendations'] = self._generate_recommendations(
                validation_report['rule_scores']
            )
            
            # Perform file analysis
            validation_report['file_analysis'] = self._analyze_exported_file(
                export_path, export_result.format_type
            )
            
        except Exception as e:
            validation_report['issues'].append(f"Validation error: {str(e)}")
            validation_report['overall_score'] = 0.2
        
        return validation_report
    
    def _validate_content_completeness(self, original: str, export_path: Path, 
                                     format_type: str) -> float:
        """Validate content completeness"""
        try:
            # Extract content from exported file
            exported_content = self._extract_content_from_file(export_path, format_type)
            
            if not exported_content:
                return 0.1
            
            # Compare content metrics
            original_words = len(original.split())
            exported_words = len(exported_content.split())
            
            # Calculate word retention ratio
            if original_words > 0:
                word_ratio = min(exported_words / original_words, 1.0)
            else:
                word_ratio = 1.0 if exported_words == 0 else 0.5
            
            # Check for major content blocks
            original_headers = len(re.findall(r'^#+\s', original, re.MULTILINE))
            exported_headers = len(re.findall(r'<h[1-6]|^#+\s', exported_content, re.MULTILINE))
            
            header_ratio = (
                min(exported_headers / original_headers, 1.0) 
                if original_headers > 0 else 1.0
            )
            
            # Combine scores
            completeness_score = (word_ratio * 0.7) + (header_ratio * 0.3)
            return max(0.0, min(1.0, completeness_score))
            
        except Exception:
            return 0.3
    
    def _validate_formatting_preservation(self, original: str, export_path: Path,
                                        format_type: str) -> float:
        """Validate formatting preservation"""
        try:
            # Count formatting elements in original
            bold_count = len(re.findall(r'\*\*[^*]+\*\*|__[^_]+__', original))
            italic_count = len(re.findall(r'\*[^*]+\*|_[^_]+_', original))
            code_count = len(re.findall(r'`[^`]+`', original))
            link_count = len(re.findall(r'\[([^\]]+)\]\([^)]+\)', original))
            
            # For HTML/EPUB, check for corresponding HTML tags
            if format_type in ['html', 'epub']:
                exported_content = self._extract_content_from_file(export_path, format_type)
                
                exported_bold = len(re.findall(r'<(strong|b)>', exported_content))
                exported_italic = len(re.findall(r'<(em|i)>', exported_content))
                exported_code = len(re.findall(r'<code>', exported_content))
                exported_links = len(re.findall(r'<a\s+[^>]*href=', exported_content))
                
                # Calculate preservation ratios
                bold_ratio = min(exported_bold / bold_count, 1.0) if bold_count > 0 else 1.0
                italic_ratio = min(exported_italic / italic_count, 1.0) if italic_count > 0 else 1.0
                code_ratio = min(exported_code / code_count, 1.0) if code_count > 0 else 1.0
                link_ratio = min(exported_links / link_count, 1.0) if link_count > 0 else 1.0
                
                return (bold_ratio + italic_ratio + code_ratio + link_ratio) / 4
            
            # For other formats, use heuristics
            return 0.7  # Assume reasonable formatting preservation
            
        except Exception:
            return 0.5
    
    def _validate_image_handling(self, original: str, export_path: Path,
                               format_type: str) -> float:
        """Validate image handling"""
        try:
            # Count images in original content
            image_refs = re.findall(r'!\[[^\]]*\]\([^)]+\)', original)
            
            if not image_refs:
                return 1.0  # No images to validate
            
            if format_type == 'html':
                exported_content = self._extract_content_from_file(export_path, format_type)
                exported_images = len(re.findall(r'<img\s+[^>]*src=', exported_content))
                
                # Check if images are preserved
                image_ratio = min(exported_images / len(image_refs), 1.0)
                return image_ratio
            
            elif format_type == 'pdf':
                # For PDF, assume images are handled if file is larger than text-only
                file_size = export_path.stat().st_size
                expected_text_size = len(original.encode('utf-8')) * 2  # Rough estimate
                
                if file_size > expected_text_size:
                    return 0.8
                else:
                    return 0.4
            
            # For other formats, provide default score
            return 0.6
            
        except Exception:
            return 0.4
    
    def _validate_reference_integrity(self, original: str, export_path: Path,
                                    format_type: str) -> float:
        """Validate reference integrity"""
        try:
            # Count different types of references
            internal_refs = len(re.findall(r'\[([^\]]+)\]\(#[^)]+\)', original))
            external_refs = len(re.findall(r'\[([^\]]+)\]\([^)#]+\)', original))
            
            total_refs = internal_refs + external_refs
            
            if total_refs == 0:
                return 1.0  # No references to validate
            
            if format_type in ['html', 'epub']:
                exported_content = self._extract_content_from_file(export_path, format_type)
                exported_links = len(re.findall(r'<a\s+[^>]*href=', exported_content))
                
                link_ratio = min(exported_links / total_refs, 1.0)
                return link_ratio
            
            # For other formats, assume reasonable reference handling
            return 0.7
            
        except Exception:
            return 0.5
    
    def _validate_file_validity(self, export_path: Path, format_type: str) -> float:
        """Validate exported file validity"""
        try:
            if not export_path.exists():
                return 0.0
            
            file_size = export_path.stat().st_size
            if file_size < 100:  # Very small file
                return 0.1
            
            # Format-specific validation
            if format_type == 'pdf':
                return self._validate_pdf_file(export_path)
            elif format_type == 'html':
                return self._validate_html_file(export_path)
            elif format_type == 'docx':
                return self._validate_docx_file(export_path)
            elif format_type == 'epub':
                return self._validate_epub_file(export_path)
            
            return 0.5  # Default for unknown formats
            
        except Exception:
            return 0.2
    
    def _validate_pdf_file(self, pdf_path: Path) -> float:
        """Validate PDF file"""
        try:
            with open(pdf_path, 'rb') as f:
                header = f.read(4)
                if header != b'%PDF':
                    return 0.2
            
            # Try to detect if it's a valid PDF structure
            with open(pdf_path, 'rb') as f:
                content = f.read()
                if b'%%EOF' not in content:
                    return 0.3
            
            return 0.8
            
        except Exception:
            return 0.2
    
    def _validate_html_file(self, html_path: Path) -> float:
        """Validate HTML file"""
        try:
            with open(html_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            score = 0.3  # Base score
            
            if '<html' in content and '</html>' in content:
                score += 0.2
            if '<head>' in content and '</head>' in content:
                score += 0.1
            if '<body>' in content and '</body>' in content:
                score += 0.2
            if '<title>' in content:
                score += 0.1
            if 'charset=' in content:
                score += 0.1
            
            return min(score, 1.0)
            
        except Exception:
            return 0.2
    
    def _validate_docx_file(self, docx_path: Path) -> float:
        """Validate DOCX file"""
        try:
            if not PYTHON_DOCX_AVAILABLE:
                return 0.5  # Can't validate without library
            
            doc = Document(docx_path)
            
            if len(doc.paragraphs) > 0:
                return 0.8
            else:
                return 0.3
                
        except Exception:
            return 0.2
    
    def _validate_epub_file(self, epub_path: Path) -> float:
        """Validate EPUB file"""
        try:
            with zipfile.ZipFile(epub_path, 'r') as z:
                files = z.namelist()
                
                score = 0.2  # Base score
                
                if 'mimetype' in files:
                    score += 0.2
                if 'META-INF/container.xml' in files:
                    score += 0.2
                if any('content.opf' in f for f in files):
                    score += 0.2
                if any('.xhtml' in f or '.html' in f for f in files):
                    score += 0.2
                
                return min(score, 1.0)
                
        except Exception:
            return 0.2
    
    def _extract_content_from_file(self, file_path: Path, format_type: str) -> str:
        """Extract text content from exported file for comparison"""
        try:
            if format_type == 'html':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Strip HTML tags for text comparison
                text = re.sub(r'<[^>]+>', ' ', content)
                return ' '.join(text.split())
            
            elif format_type == 'epub':
                if not EBOOKLIB_AVAILABLE:
                    return ""
                
                book = epub.read_epub(str(file_path))
                text_content = []
                
                for item in book.get_items():
                    if item.get_type() == ebooklib.ITEM_DOCUMENT:
                        content = item.get_content().decode('utf-8')
                        text = re.sub(r'<[^>]+>', ' ', content)
                        text_content.append(text)
                
                return ' '.join(' '.join(text_content).split())
            
            elif format_type == 'docx':
                if not PYTHON_DOCX_AVAILABLE:
                    return ""
                
                doc = Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs]
                return ' '.join(' '.join(paragraphs).split())
            
            # For other formats, return empty string
            return ""
            
        except Exception:
            return ""
    
    def _generate_recommendations(self, rule_scores: Dict[str, float]) -> List[str]:
        """Generate recommendations based on validation scores"""
        recommendations = []
        
        for rule, score in rule_scores.items():
            if score < 0.7:
                rule_info = self.validation_rules.get(rule, {})
                desc = rule_info.get('description', rule)
                
                if rule == 'content_completeness' and score < 0.7:
                    recommendations.append(
                        f"Content completeness is low ({score:.1%}). "
                        "Check if all sections are being exported properly."
                    )
                elif rule == 'formatting_preservation' and score < 0.7:
                    recommendations.append(
                        f"Formatting preservation needs improvement ({score:.1%}). "
                        "Verify that bold, italic, and other formatting is maintained."
                    )
                elif rule == 'image_handling' and score < 0.7:
                    recommendations.append(
                        f"Image handling could be improved ({score:.1%}). "
                        "Ensure images are properly embedded or referenced."
                    )
                elif rule == 'reference_integrity' and score < 0.7:
                    recommendations.append(
                        f"Reference integrity is suboptimal ({score:.1%}). "
                        "Check that internal and external links are working."
                    )
                elif rule == 'file_validity' and score < 0.7:
                    recommendations.append(
                        f"File validity concerns detected ({score:.1%}). "
                        "The exported file may have structural issues."
                    )
        
        if not recommendations:
            recommendations.append("Export quality is good. No major issues detected.")
        
        return recommendations
    
    def _analyze_exported_file(self, file_path: Path, format_type: str) -> Dict[str, Any]:
        """Perform detailed analysis of the exported file"""
        analysis = {
            'file_size_bytes': 0,
            'estimated_page_count': 0,
            'content_type': format_type,
            'creation_time': None,
            'file_health': 'unknown'
        }
        
        try:
            if file_path.exists():
                stat = file_path.stat()
                analysis['file_size_bytes'] = stat.st_size
                analysis['creation_time'] = datetime.fromtimestamp(stat.st_ctime).isoformat()
                
                # Estimate page count based on file size and type
                if format_type == 'pdf':
                    # Rough estimate: 50KB per page for PDF
                    analysis['estimated_page_count'] = max(1, stat.st_size // (50 * 1024))
                elif format_type in ['html', 'epub']:
                    # Rough estimate: 5KB per page for HTML/EPUB
                    analysis['estimated_page_count'] = max(1, stat.st_size // (5 * 1024))
                elif format_type == 'docx':
                    # Rough estimate: 10KB per page for DOCX
                    analysis['estimated_page_count'] = max(1, stat.st_size // (10 * 1024))
                
                # Basic file health check
                if stat.st_size > 0:
                    analysis['file_health'] = 'healthy'
                else:
                    analysis['file_health'] = 'empty'
            else:
                analysis['file_health'] = 'missing'
                
        except Exception as e:
            analysis['file_health'] = f'error: {str(e)}'
        
        return analysis


class ExportManager:
    """
    Main Export Manager with comprehensive functionality.
    
    This class orchestrates the entire export process, handling multiple formats,
    quality validation, and integration with the Content Management System.
    """
    
    def __init__(self, base_dir: Path, config: Optional[Dict[str, Any]] = None):
        """Initialize the Export Manager
        
        Args:
            base_dir: Base directory for export operations
            config: Optional configuration dictionary
        """
        self.base_dir = Path(base_dir)
        self.config = config or {}
        self.export_dir = self.base_dir / "export"
        self.export_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize components
        self.image_consolidator = ImageConsolidator(self.base_dir)
        self.reference_resolver = ReferenceResolver(self.base_dir)
        self.quality_validator = ExportQualityValidator()
        
        # Setup templates
        self.templates = {
            'academic': AcademicTemplate('academic', self.base_dir)
        }
        
        # Initialize exporters
        self.exporters = {}
        self._setup_exporters()
        
        # Export history
        self.export_history = []
        
    def _setup_exporters(self):
        """Setup format-specific exporters"""
        academic_template = self.templates['academic']
        
        self.exporters = {
            'pdf': PDFExporter(academic_template),
            'html': HTMLExporter(academic_template),
            'docx': DOCXExporter(academic_template),
            'epub': EPUBExporter(academic_template)
        }
    
    def export(self, content_paths: List[Union[str, Path]], output_directory: Union[str, Path],
              config: Union[Dict[str, Any], ExportConfig],
              metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Main export function that handles the complete export process.
        
        Args:
            content_paths: List of paths to content files to export
            output_directory: Directory where exported files should be saved
            config: Export configuration (dict or ExportConfig object)
            metadata: Optional metadata for the export
            
        Returns:
            Dictionary containing export results and metadata
        """
        start_time = time.time()
        
        # Ensure paths are Path objects
        content_paths = [Path(p) for p in content_paths]
        output_directory = Path(output_directory)
        output_directory.mkdir(parents=True, exist_ok=True)
        
        # Convert config if needed
        if isinstance(config, dict):
            export_config = ExportConfig(**config)
        else:
            export_config = config
        
        # Initialize result
        export_result = {
            'success': False,
            'export_id': str(uuid.uuid4()),
            'start_time': datetime.now().isoformat(),
            'format_results': {},
            'consolidated_images': {},
            'reference_warnings': [],
            'quality_reports': {},
            'summary': {},
            'errors': [],
            'warnings': []
        }
        
        try:
            # Step 1: Consolidate images if requested
            if export_config.consolidate_images:
                self._log_step("Consolidating images...")
                image_result = self.image_consolidator.consolidate_images(
                    content_paths,
                    export_config.image_sizing,
                    {
                        'quality': 85 if export_config.quality_level == 'high' else 70,
                        'dpi': 150 if export_config.quality_level == 'high' else 96
                    }
                )
                export_result['consolidated_images'] = image_result.consolidated_images
                if image_result.errors:
                    export_result['warnings'].extend(image_result.errors)
            
            # Step 2: Build reference map
            if export_config.resolve_references:
                self._log_step("Building reference map...")
                self.reference_resolver.build_reference_map(content_paths)
            
            # Step 3: Process each content file
            processed_content = []
            for content_path in content_paths:
                content = self._load_content(content_path)
                
                if export_config.resolve_references:
                    content, ref_warnings = self.reference_resolver.resolve_references(
                        content, content_path, export_result['consolidated_images']
                    )
                    export_result['reference_warnings'].extend(ref_warnings)
                
                processed_content.append(content)
            
            # Combine content
            combined_content = '\n\n'.join(processed_content)
            
            # Step 4: Export to requested formats
            formats_to_export = []
            if export_config.output_format == 'all':
                formats_to_export = ['pdf', 'html', 'docx', 'epub']
            else:
                formats_to_export = [export_config.output_format]
            
            for format_type in formats_to_export:
                if format_type in self.exporters:
                    self._log_step(f"Exporting to {format_type.upper()}...")
                    format_result = self._export_single_format(
                        combined_content,
                        output_directory,
                        format_type,
                        export_config,
                        metadata
                    )
                    export_result['format_results'][format_type] = format_result
                    
                    # Quality validation if successful
                    if format_result.success:
                        output_path = Path(format_result.output_files[0])
                        quality_report = self.quality_validator.validate_export(
                            format_result, combined_content, output_path
                        )
                        export_result['quality_reports'][format_type] = quality_report
                else:
                    export_result['errors'].append(f"Unsupported format: {format_type}")
            
            # Step 5: Generate summary
            export_result['summary'] = self._generate_summary(export_result)
            export_result['success'] = any(
                result.success for result in export_result['format_results'].values()
            )
            
            # Step 6: Update CMS records if available
            try:
                self._update_cms_records(export_result, content_paths, metadata)
            except Exception as e:
                export_result['warnings'].append(f"CMS update failed: {str(e)}")
            
        except Exception as e:
            export_result['errors'].append(f"Export process failed: {str(e)}")
        
        # Finalize result
        export_result['end_time'] = datetime.now().isoformat()
        export_result['total_time'] = time.time() - start_time
        
        # Store in history
        self.export_history.append(export_result)
        
        return export_result
    
    def _export_single_format(self, content: str, output_dir: Path, 
                             format_type: str, config: ExportConfig,
                             metadata: Dict[str, Any] = None) -> ExportResult:
        """Export content to a single format"""
        
        # Generate output filename
        base_name = metadata.get('title', 'document') if metadata else 'document'
        base_name = re.sub(r'[^\w\-_]', '_', base_name)
        
        extensions = {
            'pdf': '.pdf',
            'html': '.html',
            'docx': '.docx',
            'epub': '.epub'
        }
        
        output_path = output_dir / f"{base_name}{extensions[format_type]}"
        
        # Get exporter and export
        exporter = self.exporters[format_type]
        return exporter.export(content, output_path, config, metadata)
    
    def _load_content(self, content_path: Path) -> str:
        """Load content from file"""
        try:
            with open(content_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            return f"Error loading {content_path}: {str(e)}"
    
    def _log_step(self, message: str):
        """Log export step"""
        # In a real implementation, this would use proper logging
        print(f"[ExportManager] {message}")
    
    def _update_cms_records(self, export_result: Dict[str, Any], 
                           content_paths: List[Path],
                           metadata: Dict[str, Any] = None):
        """Update Content Management System with export records"""
        
        # This would integrate with the actual CMS
        # For now, we'll create a simple record
        cms_record = {
            'export_id': export_result['export_id'],
            'timestamp': export_result['start_time'],
            'source_files': [str(p) for p in content_paths],
            'output_formats': list(export_result['format_results'].keys()),
            'success': export_result['success'],
            'quality_scores': {
                fmt: report.get('overall_score', 0)
                for fmt, report in export_result['quality_reports'].items()
            },
            'metadata': metadata or {}
        }
        
        # Save to CMS records file
        cms_records_file = self.base_dir / "export" / "cms_records.json"
        
        try:
            if cms_records_file.exists():
                with open(cms_records_file, 'r') as f:
                    records = json.load(f)
            else:
                records = []
            
            records.append(cms_record)
            
            with open(cms_records_file, 'w') as f:
                json.dump(records, f, indent=2)
                
        except Exception as e:
            raise Exception(f"Failed to update CMS records: {str(e)}")
    
    def _generate_summary(self, export_result: Dict[str, Any]) -> Dict[str, Any]:
        """Generate export summary"""
        
        format_results = export_result['format_results']
        quality_reports = export_result['quality_reports']
        
        summary = {
            'total_formats': len(format_results),
            'successful_formats': sum(1 for r in format_results.values() if r.success),
            'failed_formats': sum(1 for r in format_results.values() if not r.success),
            'total_size_mb': sum(r.file_size_mb for r in format_results.values()),
            'average_quality_score': 0,
            'format_details': {},
            'recommendations': []
        }
        
        # Calculate average quality score
        quality_scores = [
            report.get('overall_score', 0) 
            for report in quality_reports.values()
        ]
        if quality_scores:
            summary['average_quality_score'] = sum(quality_scores) / len(quality_scores)
        
        # Format details
        for format_type, result in format_results.items():
            quality_report = quality_reports.get(format_type, {})
            summary['format_details'][format_type] = {
                'success': result.success,
                'file_size_mb': result.file_size_mb,
                'processing_time': result.processing_time,
                'quality_score': quality_report.get('overall_score', 0),
                'validation_score': result.validation_score
            }
        
        # Generate recommendations
        if summary['average_quality_score'] < 0.7:
            summary['recommendations'].append(
                "Consider reviewing export configuration for better quality."
            )
        
        if summary['failed_formats'] > 0:
            summary['recommendations'].append(
                "Some formats failed to export. Check error messages and dependencies."
            )
        
        if export_result['reference_warnings']:
            summary['recommendations'].append(
                f"Found {len(export_result['reference_warnings'])} reference issues. "
                "Review broken links and missing images."
            )
        
        return summary
    
    # Convenience methods for single-file exports
    
    def export_pdf(self, content_path: Union[str, Path], output_path: Union[str, Path],
                  template: str = 'academic', metadata: Dict[str, Any] = None) -> ExportResult:
        """Export single file to PDF"""
        config = ExportConfig(output_format='pdf', template_name=template)
        result = self.export([content_path], Path(output_path).parent, config, metadata)
        return result['format_results']['pdf']
    
    def export_html(self, content_path: Union[str, Path], output_path: Union[str, Path],
                   template: str = 'academic', metadata: Dict[str, Any] = None) -> ExportResult:
        """Export single file to HTML"""
        config = ExportConfig(output_format='html', template_name=template)
        result = self.export([content_path], Path(output_path).parent, config, metadata)
        return result['format_results']['html']
    
    def export_docx(self, content_path: Union[str, Path], output_path: Union[str, Path],
                   template: str = 'academic', metadata: Dict[str, Any] = None) -> ExportResult:
        """Export single file to DOCX"""
        config = ExportConfig(output_format='docx', template_name=template)
        result = self.export([content_path], Path(output_path).parent, config, metadata)
        return result['format_results']['docx']
    
    def export_epub(self, content_path: Union[str, Path], output_path: Union[str, Path],
                   template: str = 'academic', metadata: Dict[str, Any] = None) -> ExportResult:
        """Export single file to EPUB"""
        config = ExportConfig(output_format='epub', template_name=template)
        result = self.export([content_path], Path(output_path).parent, config, metadata)
        return result['format_results']['epub']
    
    def export_all_formats(self, content_path: Union[str, Path], output_dir: Union[str, Path],
                          template: str = 'academic', metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """Export single file to all supported formats"""
        config = ExportConfig(output_format='all', template_name=template)
        return self.export([content_path], output_dir, config, metadata)
    
    def batch_export(self, content_paths: List[Union[str, Path]], 
                    output_dir: Union[str, Path],
                    formats: List[str] = None,
                    template: str = 'academic',
                    metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """Export multiple files individually"""
        
        results = []
        formats = formats or ['pdf']
        
        for content_path in content_paths:
            for format_type in formats:
                config = ExportConfig(output_format=format_type, template_name=template)
                result = self.export([content_path], output_dir, config, metadata)
                results.append(result)
        
        return results
    
    def get_export_history(self, limit: int = 10) -> List[Dict[str, Any]]:
        """Get recent export history"""
        return self.export_history[-limit:]
    
    def get_quality_statistics(self) -> Dict[str, Any]:
        """Get quality statistics from export history"""
        if not self.export_history:
            return {'message': 'No export history available'}
        
        all_quality_scores = []
        format_stats = {}
        
        for export in self.export_history:
            for format_type, quality_report in export.get('quality_reports', {}).items():
                score = quality_report.get('overall_score', 0)
                all_quality_scores.append(score)
                
                if format_type not in format_stats:
                    format_stats[format_type] = []
                format_stats[format_type].append(score)
        
        stats = {
            'total_exports': len(self.export_history),
            'average_quality': sum(all_quality_scores) / len(all_quality_scores) if all_quality_scores else 0,
            'format_statistics': {}
        }
        
        for format_type, scores in format_stats.items():
            stats['format_statistics'][format_type] = {
                'count': len(scores),
                'average_quality': sum(scores) / len(scores),
                'min_quality': min(scores),
                'max_quality': max(scores)
            }
        
        return stats
    
    def clear_cache(self):
        """Clear export caches"""
        self.image_consolidator.image_cache.clear()
        self.reference_resolver.reference_map.clear()
        self.reference_resolver.broken_links.clear()
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported export formats"""
        return list(self.exporters.keys())
    
    def validate_dependencies(self) -> Dict[str, bool]:
        """Check availability of export dependencies"""
        return {
            'weasyprint': WEASYPRINT_AVAILABLE,
            'reportlab': REPORTLAB_AVAILABLE,
            'python_docx': PYTHON_DOCX_AVAILABLE,
            'ebooklib': EBOOKLIB_AVAILABLE,
            'markdown': MARKDOWN_AVAILABLE,
            'pillow': PILLOW_AVAILABLE
        }


# Convenience functions for quick exports

def create_export_manager(base_dir: Union[str, Path], 
                         config: Optional[Dict[str, Any]] = None) -> ExportManager:
    """Create an export manager instance"""
    return ExportManager(base_dir, config)


async def export_to_pdf(content_path: Union[str, Path], 
                       output_path: Union[str, Path],
                       metadata: Dict[str, Any] = None) -> ExportResult:
    """Quick PDF export function"""
    base_dir = Path(content_path).parent
    manager = create_export_manager(base_dir)
    return manager.export_pdf(content_path, output_path, metadata=metadata)


async def export_to_html(content_path: Union[str, Path], 
                        output_path: Union[str, Path],
                        metadata: Dict[str, Any] = None) -> ExportResult:
    """Quick HTML export function"""
    base_dir = Path(content_path).parent
    manager = create_export_manager(base_dir)
    return manager.export_html(content_path, output_path, metadata=metadata)


def export_all(content_path: Union[str, Path], 
               output_dir: Union[str, Path],
               metadata: Dict[str, Any] = None) -> Dict[str, Any]:
    """Quick export to all formats"""
    base_dir = Path(content_path).parent
    manager = create_export_manager(base_dir)
    return manager.export_all_formats(content_path, output_dir, metadata=metadata)